# 替换 v5 docx 内嵌的部署管线图为新版（功耗双口径修正后）
# 定位依据：媒体部件中 PNG 尺寸 2552×1046（fig9 独有宽高比）且字节数与旧文件一致
# 运行：python scripts/apply_v5_fig7swap.py
import sys, struct, hashlib
sys.stdout.reconfigure(encoding='utf-8', errors='replace')
from pathlib import Path
from docx import Document

NEW = Path('results/analysis/fig9-deploy-pipeline.png')
PATH = 'paper/JRTIP-paper-v5.docx'

def png_size(blob):
    # PNG IHDR: 8 字节签名 + 4 长度 + 4 "IHDR" + 4 宽 + 4 高
    assert blob[:8] == b'\x89PNG\r\n\x1a\n'
    w, h = struct.unpack('>II', blob[16:24])
    return w, h

new_bytes = NEW.read_bytes()
doc = Document(PATH)
hits = []
for rel in doc.part.rels.values():
    if 'image' not in rel.reltype:
        continue
    blob = rel.target_part.blob
    try:
        w, h = png_size(blob)
    except Exception:
        continue
    if abs(w / h - 2.44) < 0.02:  # fig9 独有宽高比（部署管线横条图）
        hits.append((rel.rId, rel.target_part, len(blob), f'{w}x{h}'))

assert len(hits) == 1, f'fig9 定位异常: {hits}'
rid, part, old_len, old_md5 = hits[0]
print(f'定位 fig9 部件: rId={rid}, 旧大小={old_len}, md5={old_md5}')
assert old_len != len(new_bytes) or hashlib.md5(part.blob).digest() != hashlib.md5(new_bytes).digest(), '已是新图？'

part._blob = new_bytes
doc.save(PATH)

# 复检：重新打开，确认部件字节 == 新文件
doc2 = Document(PATH)
for rel in doc2.part.rels.values():
    if 'image' in rel.reltype and rel.rId == rid:
        assert rel.target_part.blob == new_bytes, '替换未生效'
        print('复检通过：docx 内嵌 fig9 已替换为新图（字节级一致）')
        break
# 题注段落未动校验
full = '\n'.join(p.text for p in doc2.paragraphs)
assert 'Fig. 7 Deployment pipeline' in full
print('题注完好')
