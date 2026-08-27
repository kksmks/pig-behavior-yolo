# 生成中文补充材料 DOCX v5（表 S1：验证集分类别 AP50 [补 M5 active 格] + 表 S2：测试集分类别 AP50）
# ——与 Supplementary_Material_v5.docx 对齐（2026-08-07）
# 运行：python scripts/build_supplementary_zh_v5.py → paper/猪行为检测-补充材料-v5.docx
import sys
sys.stdout.reconfigure(encoding='utf-8', errors='replace')
from docx import Document
from docx.shared import Pt

doc = Document()
doc.styles['Normal'].font.name = 'SimSun'
doc.styles['Normal'].font.size = Pt(11)

doc.add_heading('补充材料（在线资源 1 / Online Resource 1）', level=1)
doc.add_paragraph('面向边缘设备的群养猪多行为实时检测：类别不均衡感知采样与轻量化 FasterNet 主干 —— 补充材料')
doc.add_paragraph('Journal of Real-Time Image Processing · 张展 · 勾林林 · '
                        '通讯作者邮箱：851709772@qq.com')

cap = doc.add_paragraph('表 S1 验证集分类别 AP50。独立测试集上的关键稀有类结果：'
                        'active 0.526→0.639（基线→M4，+11.3 点）；sitting 0.534→0.571；fight 0.840→0.871。')
cap.runs[0].font.size = Pt(9)

headers = ['类别', '基线', 'M3（FasterNet）', 'M4（采样）', 'M5（组合）', 'M4−基线']
rows = [
    ['active', '0.459', '0.372', '0.552', '0.452', '+9.3'],
    ['drink', '0.408', '0.440', '0.459', '0.430', '+5.1'],
    ['eat', '0.436', '0.437', '0.447', '0.449', '+1.1'],
    ['fight', '0.858', '0.828', '0.842', '0.786', '-1.6'],
    ['investigating', '0.581', '0.571', '0.570', '0.577', '-1.1'],
    ['lying', '0.783', '0.764', '0.766', '0.743', '-1.7'],
    ['nose-to-nose', '0.686', '0.655', '0.622', '0.710', '-6.4'],
    ['sitting', '0.403', '0.233', '0.423', '0.349', '+2.0'],
    ['standing', '0.438', '0.429', '0.449', '0.445', '+1.1'],
    ['walk', '0.687', '0.646', '0.686', '0.667', '-0.1'],
    ['mAP50（全部）', '0.573', '0.537', '0.582', '0.561', '+0.9'],
]
t = doc.add_table(rows=1 + len(rows), cols=len(headers))
t.style = 'Light Grid Accent 1'
for j, htxt in enumerate(headers):
    t.rows[0].cells[j].text = htxt
for i, row in enumerate(rows):
    for j, v in enumerate(row):
        t.rows[i + 1].cells[j].text = v

# ===== 表 S2：测试集分类别 AP50（2026-08-07 官方协议复核） =====
cap2 = doc.add_paragraph('表 S2 独立测试集上的分类别 AP50（基线与两个采用变体）。稀有类格局与验证集'
                         '（表 S1）一致：M4 的增益集中于 active、sitting 与 fight；M5 保留部分稀有类增益，'
                         '同时参数量少 4.4%。所有数值为单次实验（种子 0）；整体指标的三种子均值见正文 4.5 节。'
                         'M5 数值在官方验证协议下复测（整体 mAP50 0.5933，与云端所报 0.5932 一致）。')
cap2.runs[0].font.size = Pt(9)

headers2 = ['类别', '基线', 'M4（采样）', 'M5（组合）']
rows2 = [
    ['active', '0.526', '0.639', '0.533'],
    ['drink', '0.389', '0.373', '0.408'],
    ['eat', '0.445', '0.418', '0.429'],
    ['fight', '0.840', '0.871', '0.849'],
    ['investigating', '0.552', '0.562', '0.588'],
    ['lying', '0.741', '0.737', '0.722'],
    ['nose-to-nose', '0.783', '0.740', '0.749'],
    ['sitting', '0.534', '0.571', '0.568'],
    ['standing', '0.487', '0.468', '0.475'],
    ['walk', '0.667', '0.655', '0.612'],
    ['mAP50（全部）', '0.596', '0.604', '0.593'],
]
t2 = doc.add_table(rows=1 + len(rows2), cols=len(headers2))
t2.style = 'Light Grid Accent 1'
for j, htxt in enumerate(headers2):
    t2.rows[0].cells[j].text = htxt
for i, row in enumerate(rows2):
    for j, v in enumerate(row):
        t2.rows[i + 1].cells[j].text = v

# ===== 图 S1–S3（2026-08-18 新增：PR 曲线 / 延迟分解 / 持续运行时序） =====
from pathlib import Path
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

FIGS = [
    ('results/analysis/figS1-pr-curves.png',
     '图 S1 独立测试集上的精确率-召回率（PR）曲线（四个消融模型）：(a) 基线 YOLOv11n（mAP50 0.596）、'
     '(b) M3 仅 FasterNet（0.569）、(c) M4 仅采样（0.604）、(d) 最终 M5（0.593）。'
     '细线为各行为类别曲线，粗蓝线为全类平均'),
    ('results/analysis/figS2-latency-breakdown.png',
     '图 S2 Jetson Nano 延迟预算。(a) 引擎侧分解（TensorRT 实测）：两种输入尺寸下 GPU 计算均占绝对主导，'
     'H2D/D2H 传输可忽略。(b) 纯 Python 管线的 CPU 侧前/后处理开销（对数轴）：朴素的解码与缩放'
     '可能超过 GPU 计算，生产管线应以 C 或 CUDA 实现前/后处理'),
    ('results/analysis/figS3-thermal-timeline.png',
     '图 S3 Jetson Nano 持续运行时序（M5 @640 FP16，连续 12,000 次推理、10.7 分钟）。'
     '吞吐保持 20.0 FPS（第 99 百分位延迟 51.3 ms）；GPU 温度峰值 55.5 °C，全程零热节流、无降频'),
]
for path, caption in FIGS:
    assert Path(path).exists(), f'缺图: {path}'
    doc.add_picture(path, width=Inches(6.2))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].font.size = Pt(9)

out = 'paper/猪行为检测-补充材料-v5.docx'
doc.save(out)
print('已生成:', out)
