# 把图 S1–S3 追加到英文补充材料（Supplementary_Material_v5.docx 尾部，Springer 图题式）
# 运行：python scripts/apply_supp_figures.py（需 figS1/figS2/figS3 已就位）
import sys
sys.stdout.reconfigure(encoding='utf-8', errors='replace')
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

PATH = 'paper/Supplementary_Material_v5.docx'
FIGS = [
    ('results/analysis/figS1-pr-curves.png',
     'Fig. S1 Precision–recall curves on the held-out test set for the four ablation models: '
     '(a) baseline YOLOv11n (mAP50 0.596), (b) M3 FasterNet-only (0.569), (c) M4 sampling-only '
     '(0.604), (d) final M5 (0.593). Behaviour-class curves are shown thin; the thick blue curve '
     'is the all-class average'),
    ('results/analysis/figS2-latency-breakdown.png',
     'Fig. S2 Latency budget on the Jetson Nano. (a) Engine-side decomposition (TensorRT '
     'measurement): GPU compute dominates at both input sizes; H2D/D2H transfers are negligible. '
     '(b) CPU-side pre/post-processing costs in a pure-Python pipeline (log scale): naive decode '
     'and resize can exceed GPU compute, so production pipelines should implement pre/'
     'post-processing in C or CUDA'),
    ('results/analysis/figS3-thermal-timeline.png',
     'Fig. S3 Sustained-operation timeline on the Jetson Nano over 12,000 consecutive inferences '
     '(10.7 min, M5 at 640×640 FP16). Throughput holds at 20.0 FPS (99th-percentile latency '
     '51.3 ms); GPU temperature peaks at 55.5 °C with zero thermal-throttling events and no clock '
     'drop'),
]

doc = Document(PATH)
n0 = len(doc.paragraphs)
for path, caption in FIGS:
    assert Path(path).exists(), f'缺图: {path}'
    doc.add_picture(path, width=Inches(6.2))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].font.size = Pt(9)
doc.save(PATH)

# 复检
doc2 = Document(PATH)
imgs = [r for r in doc2.part.rels.values() if 'image' in r.reltype]
full = '\n'.join(p.text for p in doc2.paragraphs)
assert len(imgs) == 3, f'图片数异常: {len(imgs)}'
for probe in ['Fig. S1', 'Fig. S2', 'Fig. S3']:
    assert probe in full, probe
print(f'OK: 3 图已追加（段落 {n0}→{len(doc2.paragraphs)}），题注齐')
