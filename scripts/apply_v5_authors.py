# v5 注入作者信息（独作双人共一版：Zhan Zhang & Linlin Gou，六盘水师范学院）
# 运行：python scripts/apply_v5_authors.py → 原地更新 paper/JRTIP-paper-v5.docx
import sys
sys.stdout.reconfigure(encoding='utf-8', errors='replace')
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

PATH = 'paper/JRTIP-paper-v5.docx'
doc = Document(PATH)
paras = doc.paragraphs

# ① 标题页作者块（para 1 占位行 → 多行作者块）
i = next(i for i, p in enumerate(paras) if p.text.strip() == '[Authors / affiliations to be filled]')
p = paras[i]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for r in list(p.runs):
    r.text = ''
lines = [
    ('Zhan Zhang†* and Linlin Gou†', False),
    ('School of Biological Science and Technology, Liupanshui Normal University, '
     'Liupanshui 553000, China', True),
    ('† These authors contributed equally to this work.', True),
    ('*Corresponding author: Zhan Zhang, e-mail: 851709772@qq.com', True),
]
first = True
for text, small in lines:
    run = p.add_run(text) if first else p.add_run()
    if not first:
        run.add_break()
        run.add_text(text)
    run.font.size = Pt(10 if small else 12)
    run.italic = False
    first = False
print('OK 标题页作者块')

# ② Author Contributions
i = next(i for i, p in enumerate(paras) if p.text.startswith('Author Contributions:'))
p = paras[i]
for r in list(p.runs):
    r.text = ''
p.runs[0].text = ('Author Contributions: Z.Z. and L.G. contributed equally to this work: '
                  'conceptualization, methodology, software, validation, formal analysis, '
                  'investigation, data curation, writing—original draft, writing—review and '
                  'editing, and visualization.')
print('OK Author Contributions')

# ③ Funding
i = next(i for i, p in enumerate(paras) if p.text.startswith('Funding:'))
p = paras[i]
for r in list(p.runs):
    r.text = ''
p.runs[0].text = 'Funding: The authors received no specific funding for this work.'
print('OK Funding')

# ④ 作者 bio（JRTIP 硬性要求，追加到参考文献之后）
bios = [
    'Zhan Zhang is a third-year undergraduate student majoring in Animal Science at the School of '
    'Biological Science and Technology, Liupanshui Normal University, Liupanshui, China. Research '
    'interests include computer vision, edge computing, and their applications to precision '
    'livestock farming. This is a first peer-reviewed publication.',
    'Linlin Gou is an undergraduate student at the School of Biological Science and Technology, '
    'Liupanshui Normal University, Liupanshui, China. Research interests include animal science '
    'and applied machine learning.',
]
for b in bios:
    par = doc.add_paragraph()
    run = par.add_run(b)
    run.font.size = Pt(9)
print('OK bio ×2')

doc.save(PATH)

# 复检
full = '\n'.join(p.text for p in Document(PATH).paragraphs)
for probe in ['Zhan Zhang†* and Linlin Gou†', 'contributed equally', '553000',
              '851709772@qq.com', 'contributed equally to this work: conceptualization',
              'no specific funding', 'third-year undergraduate student majoring in Animal Science']:
    assert probe in full, f'复检失败: {probe[:40]}'
assert '[Authors / affiliations to be filled]' not in full
assert '[to be completed by the authors]' not in full
print('复检通过：作者块/贡献/基金/bio 全部落位，占位符零残留')
