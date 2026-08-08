#!/usr/bin/env python3
"""修复 Supplementary_Material_v5.docx 的表顺序：S2 标题+表应位于 S1 表之后。

python-docx 的 add_paragraph/add_table 在该文档上把新内容插到了 S1 表之前，
导致 S1 标题与 S1 表分离。本脚本从原始 Supplementary_Material.docx 重新生成，
用显式 XML addnext 保证顺序：…→ S1 标题 → S1 表 → S2 标题 → S2 表。

运行：python scripts/fix_supp_v5_order.py
"""
import copy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

PAPER = Path(r'E:/pig-behavior-yolo/paper')
SUPP_SRC = PAPER / 'Supplementary_Material.docx'
SUPP_DST = PAPER / 'Supplementary_Material_v5.docx'

S2_ROWS = [
    ("active",        "0.526", "0.639", "0.533"),
    ("drink",         "0.389", "0.373", "0.408"),
    ("eat",           "0.445", "0.418", "0.429"),
    ("fight",         "0.840", "0.871", "0.849"),
    ("investigating", "0.552", "0.562", "0.588"),
    ("lying",         "0.741", "0.737", "0.722"),
    ("nose-to-nose",  "0.783", "0.740", "0.749"),
    ("sitting",       "0.534", "0.571", "0.568"),
    ("standing",      "0.487", "0.468", "0.475"),
    ("walk",          "0.667", "0.655", "0.612"),
    ("mAP50 (all)",   "0.596", "0.604", "0.593"),
]
S2_CAPTION = ("Table S2. Per-class AP50 on the held-out test set for the baseline and the two "
              "adopted variants. The rare-class pattern mirrors the validation set (Table S1): "
              "the gains of M4 concentrate on active, sitting and fight; M5 keeps part of the "
              "rare-class gain while carrying 4.4% fewer parameters. M5 values were re-measured "
              "under the official validation protocol (overall mAP50 0.5933, matching the "
              "cloud-reported 0.5932).")

doc = Document(str(SUPP_SRC))
s1 = doc.tables[0]

# 1) 补 S1 的 M5 active 格
header = [c.text.strip() for c in s1.rows[0].cells]
m5_col = header.index('M5 (combined)')
for row in s1.rows[1:]:
    if row.cells[0].text.strip() == 'active':
        assert row.cells[m5_col].text.strip() in ('—', '-', ''), row.cells[m5_col].text
        row.cells[m5_col].text = '0.452'
        print('[OK] S1 M5 active → 0.452')

# 2) 在文档末尾正常创建 S2 标题 + 表（先借 python-docx 生成 XML）
cap_para = doc.add_paragraph(S2_CAPTION)
s2 = doc.add_table(rows=1 + len(S2_ROWS), cols=4)
try:
    s2.style = s1.style
except Exception as e:
    print('[WARN] 样式复制失败:', e)
for j, h in enumerate(['Class', 'Baseline', 'M4 (sampling)', 'M5 (combined)']):
    s2.rows[0].cells[j].text = h
for i, row in enumerate(S2_ROWS):
    for j, v in enumerate(row):
        s2.rows[i + 1].cells[j].text = v

# 3) 显式移动 XML：S1 表 → S2 标题 → S2 表
s1_tbl_el = s1._element
cap_el = cap_para._element
s2_tbl_el = s2._element
# 先从当前位置摘下
cap_el.getparent().remove(cap_el)
s2_tbl_el.getparent().remove(s2_tbl_el)
# 依次插到 S1 表之后
s1_tbl_el.addnext(cap_el)
cap_el.addnext(s2_tbl_el)

doc.save(str(SUPP_DST))
print(f'[OK] 已重排并保存 → {SUPP_DST}')

# 4) 验证文档顺序
doc2 = Document(str(SUPP_DST))
from docx.table import Table
from docx.text.paragraph import Paragraph
for child in doc2.element.body.iterchildren():
    tag = child.tag.split('}')[1]
    if tag == 'p':
        t = Paragraph(child, doc2).text.strip()
        if t:
            print('P :', t[:70])
    elif tag == 'tbl':
        tb = Table(child, doc2)
        print(f'TBL: {len(tb.rows)}x{len(tb.columns)} | 首行: ' +
              ' | '.join(c.text.strip() for c in tb.rows[0].cells))
