# v5 部署章补强：功耗双口径修正（~5W→7W总输入/5W计算轨）+ 延迟三段分解/热节流/功耗实测段落
# 依据：results/EXPERIMENT_LOG.md 2026-08-18 两行实测（trtexec 分解 + tegrastats/INA3221）
# 运行：python scripts/apply_v5_breakdown.py  → 原地更新 paper/JRTIP-paper-v5.docx
import sys
sys.stdout.reconfigure(encoding='utf-8', errors='replace')
from docx import Document

PATH = 'paper/JRTIP-paper-v5.docx'
doc = Document(PATH)
paras = doc.paragraphs

def swap_para_text(idx, old, new):
    p = paras[idx]
    txt = p.text
    assert old in txt, f'锚点未命中 para {idx}: {old[:60]}'
    new_txt = txt.replace(old, new)
    # 清空 runs 后写入（首个 run 承载全部文本，保持段落级样式）
    for r in list(p.runs):
        r.text = ''
    if p.runs:
        p.runs[0].text = new_txt
    else:
        p.add_run(new_txt)
    print(f'OK para {idx}: ...{new[:70]}')

# ① 摘要：at roughly 5 W total power → 双口径
i3 = next(i for i, p in enumerate(paras) if p.text.startswith('Automated monitoring'))
swap_para_text(i3, 'at roughly 5 W total power',
               'at roughly 7 W total board power (≈5 W across the compute rails)')

# ② 引言贡献（3）：≈5 W → ≈7 W
i12 = next(i for i, p in enumerate(paras) if p.text.startswith('(3) A deployment validation'))
swap_para_text(i12, '(19.7–33.3 FPS at ≈5 W)', '(19.7–33.3 FPS at ≈7 W board power)')

# ③ 部署段：approximately 5 W under load → 7 W + 指引下文
i68 = next(i for i, p in enumerate(paras) if p.text.startswith('A detector that only runs'))
swap_para_text(i68, 'with the whole board drawing approximately 5 W under load',
               'with the whole board drawing ≈7 W under sustained load (rail-level breakdown below)')

# ④ 结论：at approximately 5 W → 7 W total
i86 = next(i for i, p in enumerate(paras) if p.text.startswith('This paper presented'))
swap_para_text(i86, 'at approximately 5 W', 'at approximately 7 W total power')

# ⑤ 在 "three practical findings" 段之后、第 6 章之前插入三段分解段
NEW = ('Three further measurements delimit the real-time claim. First, the latency budget '
       'decomposes cleanly at 640×640: host-to-device transfer costs 0.48 ms, GPU compute '
       '49.6 ms, and device-to-host 0.06 ms—and the engine is therefore not the bottleneck to '
       'watch. The CPU side is: a pure-Python pre/post pipeline would exceed the GPU time '
       '(JPEG decode and resize cost 17.1 ms from a 640×640 source and 107.9 ms from a 1080p '
       'source; normalize and layout conversion cost 6.8 ms in optimized C; NMS over ≈370 '
       'candidates costs 0.5 ms), and TensorRT enqueue overhead adds 9.5–10.4 ms, so '
       'production pipelines should implement pre/post-processing in C or CUDA. Second, '
       'sustained operation is stable: over 12,000 consecutive inferences (10.7 min), '
       'throughput held at 20.0 FPS with a 99th-percentile latency of 51.3 ms, GPU '
       'temperature peaked at 55.5 °C, and no thermal-throttling event or clock drop '
       'occurred. Third, rail-level measurements (INA3221) show 3.4 W at idle and ≈7.0 W '
       'total board input under sustained load, of which ≈5.0 W falls on the GPU and CPU '
       'rails—this is the figure headlined above.')
i_gen = next(i for i, p in enumerate(paras) if p.text.startswith('6. Generalization Analysis'))
anchor = paras[i_gen]
new_p = anchor.insert_paragraph_before(NEW)
# 与被插入位置相同的正文样式
new_p.style = paras[i68].style
print('OK 插入三段分解段（位于 Generalization 之前）')

doc.save(PATH)

# 复检：五处全部落位、旧措辞零残留
doc2 = Document(PATH)
full = '\n'.join(p.text for p in doc2.paragraphs)
for probe in ['7 W total board power (≈5 W across the compute rails)',
              'at ≈7 W board power)', '≈7 W under sustained load (rail-level breakdown below)',
              'at approximately 7 W total power',
              '12,000 consecutive inferences', '55.5 °C', 'INA3221', '9.5–10.4 ms']:
    assert probe in full, f'复检失败: {probe[:50]}'
assert 'at roughly 5 W' not in full and 'at approximately 5 W' not in full and 'at ≈5 W' not in full
print('复检通过：5 处修改 + 1 段新增全部落位，旧 5W 措辞零残留')
