#!/usr/bin/env python3
"""论文 v3 → v4 自动文本替换脚本。

功能：
  - 自动替换所有安全的 Fig/Table 编号引用
  - 保存为新文件 JRTIP-paper-v4-partial.docx
  - 输出还需要手动做的清单

用法：
  python scripts/auto_ref_numbering.py

注意：
  - 只改文本引用，不改图表物理内容
  - 改完后务必人工通读检查
"""

import sys
from pathlib import Path

try:
    import docx
except ImportError:
    print("[错误] 需要 python-docx: pip install python-docx")
    sys.exit(1)

SRC = Path(r"E:/pig-behavior-yolo/paper/JRTIP-paper-v3.docx")
DST = Path(r"E:/pig-behavior-yolo/paper/JRTIP-paper-v4-partial.docx")

# ---------------------------------------------------------------------------
# 替换规则（按安全顺序：先替换不会被后续规则误伤的）
# ---------------------------------------------------------------------------

# Fig 替换顺序：原Fig.8→Fig.6 必须先做，否则Fig.10/11→Fig.8会误伤
FIG_RULES = [
    ("Fig. 8", "Fig. 6"),   # 原Fig.8 混淆矩阵 → 新Fig.6
    ("Fig. 9", "Fig. 7"),   # 原Fig.9 部署管线 → 新Fig.7
    ("Fig. 10", "Fig. 8"),  # 原Fig.10 泛化柱状 → 新Fig.8
    ("Fig. 11", "Fig. 8"),  # 原Fig.11 泛化样例 → 新Fig.8
    ("Fig. 5", "Fig. 4"),   # 原Fig.5 帕累托 → 新Fig.4
]

# Table 替换（只做无连环风险的）
TABLE_RULES = [
    ("Table 10", "Table 5"),  # 原T10 跨数据集 → 新T5
    ("Table 9", "Table 5"),   # 原T9 序列不相交 → 新T5
    ("Table 7", "Table 6"),   # 原T7 效率 → 新T6
]

# 需要删除或特殊处理的引用标记（脚本会在这些位置打印警告）
DELETE_FIG_REFS = ["Fig. 4"]   # 原Fig.4 柱状图要删除
DELETE_TABLE_REFS = ["Table 2", "Table 8"]  # T2并入T1, T8并入正文

SPECIAL_HANDLE = {
    "Fig. 6": "[保留，但内容需与Fig.7合并成新Fig.5]",
    "Fig. 7": "[删除，并入Fig.5]",
    "Table 1": "[保留，但需与Table 2合并内容]",
    "Table 3": "[保留，但编号需改为Table 2]",
    "Table 4": "[保留，但编号需改为Table 3]",
    "Table 5": "[保留，但编号需改为Table 4]",
    "Table 6": "[移到补充材料，正文只保留关键数字]",
}


def apply_replacements(doc, rules, label):
    changes = []
    for para in doc.paragraphs:
        original = para.text
        new_text = original
        for old, new in rules:
            if old in new_text:
                new_text = new_text.replace(old, new)
        if new_text != original:
            para.text = new_text
            changes.append((original[:80], new_text[:80]))
    print(f"\n[{label}] 完成，共修改 {len(changes)} 个段落")
    for orig, new in changes[:5]:
        print(f"  {orig[:60]}... → {new[:60]}...")
    if len(changes) > 5:
        print(f"  ... 等共 {len(changes)} 处")
    return changes


def flag_special(doc):
    print("\n[特殊处理标记] 以下引用需要手动处理：")
    found = []
    for i, para in enumerate(doc.paragraphs):
        t = para.text
        for ref, note in SPECIAL_HANDLE.items():
            if ref in t:
                found.append((i, ref, t[:100], note))
        for ref in DELETE_FIG_REFS + DELETE_TABLE_REFS:
            if ref in t:
                found.append((i, ref, t[:100], "[需要删除或改写引用]"))
    
    for idx, ref, text, note in found:
        print(f"  段落{idx}: [{ref}] {note}")
        print(f"    原文: {text}...")
    return found


def main():
    print(f"读取: {SRC}")
    doc = docx.Document(str(SRC))
    print(f"段落数: {len(doc.paragraphs)}, 表格数: {len(doc.tables)}")

    # 1. Fig 替换
    fig_changes = apply_replacements(doc, FIG_RULES, "Fig 编号替换")

    # 2. Table 替换
    table_changes = apply_replacements(doc, TABLE_RULES, "Table 编号替换")

    # 3. 标记特殊处理
    special = flag_special(doc)

    # 4. 保存
    DST.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(DST))
    print(f"\n[OK] 已保存: {DST}")

    # 5. 输出手动清单
    manual = []
    manual.append("\n" + "="*60)
    manual.append("还需要你手动做的事（在 Word 里打开 v4-partial.docx 操作）：")
    manual.append("="*60)
    manual.append("")
    manual.append("【图操作】")
    manual.append("1. 删除原 Fig.4（每类AP50柱状图）及其图题段落")
    manual.append("2. 删除原 Fig.7（难例图），把内容并入 Fig.6 旁边")
    manual.append("3. 用 PPT/Photoshop 把原 Fig.6 + Fig.7 拼成一张左右子图 → 新 Fig.5")
    manual.append("4. 用 PPT/Photoshop 把原 Fig.10 + Fig.11 拼成一张上下子图 → 新 Fig.8")
    manual.append("5. 重排全文图号引用（脚本已改大部分，但需人工检查）")
    manual.append("")
    manual.append("【表操作】")
    manual.append("6. 合并 Table 1 + Table 2 为新 Table 1（四列：定义+实例数+图像数）")
    manual.append("7. 删除 Table 8（部署表），把两行数据直接写入正文第5节")
    manual.append("8. 把 Table 6（每类AP详细数字）移到新建文件 Supplementary_Material.docx")
    manual.append("9. Table 3→改号为Table 2, Table 4→改号为Table 3, Table 5→改号为Table 4")
    manual.append("   （脚本未改这3个，因为连环替换风险，建议手动 Ctrl+H）")
    manual.append("")
    manual.append("【Ctrl+H 全局替换建议】")
    manual.append("  查找: \"Table 3\" → 替换: \"Table 2\"")
    manual.append("  查找: \"Table 4\" → 替换: \"Table 3\"")  
    manual.append("  查找: \"Table 5\" → 替换: \"Table 4\"")
    manual.append("  （注意：只在正文中替换，不要替换图题里的）")
    manual.append("")
    manual.append("【最后检查】")
    manual.append("10. 通读全文，确认所有 Fig.1-8 和 Table.1-6 引用无跳号/重复")
    manual.append("11. 另存为 JRTIP-paper-v4.docx")
    manual.append("="*60)

    report = "\n".join(manual)
    print(report)

    # 保存清单到文件
    report_path = DST.parent / "v4-manual-todo.txt"
    report_path.write_text(report, encoding="utf-8")
    print(f"\n[OK] 手动清单已保存: {report_path}")


if __name__ == "__main__":
    main()
