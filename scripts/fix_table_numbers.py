#!/usr/bin/env python3
"""补充：Table 3→2, 4→3, 5→4 的连环替换（用占位符避免误伤）。

用法：
  python scripts/fix_table_numbers.py

输入：JRTIP-paper-v4-partial.docx
输出：JRTIP-paper-v4.docx
"""

from pathlib import Path
import docx

SRC = Path(r"E:/pig-behavior-yolo/paper/JRTIP-paper-v4-partial.docx")
DST = Path(r"E:/pig-behavior-yolo/paper/JRTIP-paper-v4.docx")

def main():
    doc = docx.Document(str(SRC))
    print(f"读取: {SRC} ({len(doc.paragraphs)} 段落)")

    # Step 1: 占位符替换（避免连环）
    placeholder_rules = [
        ("Table 5", "__TABLE5_PLACEHOLDER__"),
        ("Table 4", "__TABLE4_PLACEHOLDER__"),
        ("Table 3", "__TABLE3_PLACEHOLDER__"),
    ]

    count = 0
    for para in doc.paragraphs:
        original = para.text
        new_text = original
        for old, new in placeholder_rules:
            if old in new_text:
                new_text = new_text.replace(old, new)
        if new_text != original:
            para.text = new_text
            count += 1
    print(f"[Step 1] 占位符替换: {count} 个段落")

    # Step 2: 占位符 → 最终编号
    final_rules = [
        ("__TABLE5_PLACEHOLDER__", "Table 4"),
        ("__TABLE4_PLACEHOLDER__", "Table 3"),
        ("__TABLE3_PLACEHOLDER__", "Table 2"),
    ]

    count = 0
    for para in doc.paragraphs:
        original = para.text
        new_text = original
        for old, new in final_rules:
            if old in new_text:
                new_text = new_text.replace(old, new)
        if new_text != original:
            para.text = new_text
            count += 1
    print(f"[Step 2] 最终编号替换: {count} 个段落")

    doc.save(str(DST))
    print(f"\n[OK] 已保存: {DST}")
    print("现在 Fig/Table 编号已全部自动化，你只需要做图表物理操作（删图/合并图/合并表）。")


if __name__ == "__main__":
    main()
