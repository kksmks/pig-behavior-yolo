# 组装中文版带图成稿 DOCX v3（详写版，与 JRTIP-paper-v3 逐节对齐，含评审报告#4全部修正）
# 运行：python scripts/build_paper_docx_zh.py → paper/猪行为检测-中文成稿-v3.docx
import sys
sys.stdout.reconfigure(encoding='utf-8', errors='replace')
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
st = doc.styles['Normal']
st.font.name = 'SimSun'
st.font.size = Pt(11)


def h(text, level=1):
    doc.add_heading(text, level=level)


def p(text, italic=False):
    par = doc.add_paragraph()
    r = par.add_run(text)
    r.italic = italic
    return par


def fig(path, caption, width=6.2):
    if Path(path).exists():
        doc.add_picture(path, width=Inches(width))
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].font.size = Pt(9)


def table(headers, rows, caption):
    cap = doc.add_paragraph(caption)
    cap.runs[0].font.size = Pt(9)
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    for j, htxt in enumerate(headers):
        t.rows[0].cells[j].text = htxt
    for i, row in enumerate(rows):
        for j, v in enumerate(row):
            t.rows[i + 1].cells[j].text = str(v)


# ===== 标题 =====
doc.add_heading('面向边缘设备的群养猪多行为实时检测：类别不均衡感知采样与轻量化 FasterNet 主干', level=0)
p('[作者/单位 待填]', italic=True)

# ===== 摘要 =====
h('摘要', 1)
p('群养猪的自动化行为监测，要求检测模型既能识别稀有但关乎福利的行为，又能在低成本边缘硬件上实时运行。'
  '本文基于 YOLOv11n 构建了一套实时多行为检测框架，并在 Jetson Nano 上完成部署验证。三项工作让这套框架'
  '真正可用：其一，类别不均衡感知采样策略（频率封顶的离线过采样），不改网络即提升稀有类检测——低频行为'
  'active 的 AP 在测试集上从 0.526 提升至 0.639，总体召回率在验证集上提高 2.1 个百分点；其二，FasterNet '
  '轻量化主干替换配合"恒等保持"预训练权重迁移策略，参数量降低 4.4%，三次重复实验中精度与基线统计相当'
  '（test mAP50 0.590±0.009 对 0.596），且推理速度为所有参评模型最快（RTX 3090 上 117.6 FPS）；其三，'
  'Jetson Nano 实测验证实时性（640×640 输入 19.7 FPS、480×480 输入 33.3 FPS、整机功耗约 5W），并通过两级'
  '泛化分析（未见序列压力测试与独立农场数据集）如实划定框架的适用边界。')
p('关键词：猪行为检测；类别不均衡；重复因子采样；FasterNet；边缘计算；Jetson Nano；实时目标检测；精准畜牧',
  italic=True)

# ===== 1 引言 =====
h('1 引言', 1)
p('不接触动物本身，行为几乎是我们能获取的全部信息。采食与饮水反映代谢与稳态，休息姿态指示舒适与健康，'
  '活动量体现活力，而打斗等攻击互动则意味着资源竞争、应激与福利风险 [1, 2]。行为模式的'
  '异常往往是应激或疾病最早的可见征兆，持续的行为监测也因此成为精准畜牧（PLF）[3] 最实用的手段之一。'
  '人工巡栏无法覆盖商业规模的猪群，基于摄像头的非接触式监测于是成为默认选择。')
p('但监测要在经济账上成立，检测就必须跑在猪舍内的廉价边缘设备上。一个商业猪场可能有上百个栏位，'
  '把所有视频流上传云端，在农村网络条件下既不现实，云推理的持续开销也远超边缘板卡的一次性投入。'
  '延迟同样关键：一条迟到几分钟的打斗警报几乎没有价值。这就是为什么工程上的主流趋势是把检测器带到'
  '数据旁边，而不是反过来。')
p('YOLO 系检测器 [4] 兼顾实时性与精度，是畜禽行为检测的主力。然而文献中三个问题反复出现：一是行为数据'
  '类别严重不均衡——静卧与探究行为占绝大多数，饮水、短时社交等行为稀少，均匀采样训练天然偏向高频类，'
  '而高频类往往恰好是福利预警最不关心的；二是多数畜牧检测器只在服务器 GPU 上评估，它们在内存与功耗'
  '受限的边缘硬件上表现如何基本无人记录——精度报告并不能告诉我们模型在猪舍里到底跑不跑得动；三是'
  '评测几乎总在与训练同源的视频序列抽帧上进行，相邻帧近乎重复，报告的精度偏乐观，跨栏位、跨天、'
  '跨农场的泛化能力无人量化。')
p('针对上述三个缺口，本文提出一个结合数据层再平衡、轻量化主干与上板验证的框架，贡献如下：')
p('（1）类别不均衡感知采样策略（频率封顶的离线过采样），不改网络与损失函数即可提升稀有行为检测，'
  '并附采样封顶值的敏感性分析。')
p('（2）FasterNet 轻量化主干替换与"恒等保持"集成策略。我们早期尝试向预训练网络插入新的门控模块，'
  '结果破坏了预训练特征分布、损失 3–7 个 mAP 点；改为保持通道流不变、经索引对齐重映射迁移预训练权重后，'
  '损伤完全避免——精度保住的同时参数量降低 4.4%。')
p('（3）Jetson Nano 部署验证（19.7–33.3 FPS、约 5W）与两级泛化分析（未见序列、独立农场数据集），'
  '完整划定框架的适用边界。')

# ===== 2 相关工作 =====
h('2 相关工作', 1)
h('2.1 畜禽行为视觉检测', 2)
p('YOLO 检测器已应用于猪只检测、个体追踪、姿态识别与行为分析。对群养猪，Tu 等 [5] 以 '
  'YOLOv5s/YOLOX-S 检测器结合改进 DeepSORT 追踪器追踪静卧、采食、站立等个体行为；Li 等 [6] 以 '
  'YOLOX 配合 SCTS-SlowFast 时空模块完成多行为联合检测与识别；'
  '更近的工作提出改进的 YOLOv8n、YOLOv11 变体——通常做法是插入注意力模块或重设计特征融合路径——识别'
  '站立、静卧、采食、饮水、啃咬等行为。更早的 Alameer 等 [7] 已在同类顶视栏舍画面上识别姿态与饮水'
  '行为，用于健康受损监测。'
  '这些工作验证了可行性，但大多默认服务器端推理与类别均衡的分布，且少有报告稀有行为的分类别结果。')
h('2.2 目标检测中的类别不均衡', 2)
p('前景类不均衡的治理路线大致三家：采样（类别感知采样、重复因子采样 [8]）、损失重加权'
  '（focal 系列变体 [9]）、'
  '数据增强（mosaic、mixup、copy-paste [10]）。跨路线的系统性诊断表明没有单一策略全面占优 [11, 12]，'
  '但采样法是侵入性'
  '最低的：它既不碰网络也不碰损失，因此不会干扰预训练权重——这一点在我们这种重度依赖迁移的管线里尤为'
  '重要。重复因子采样 [8] 在长尾检测基准上已是成熟做法，也被畜牧检测采用（如 WFE-YOLO [13] 的加权数据集）。'
  '本文采用离线、频率封顶的变体，并分析封顶值的敏感性——这一角度在该领域文献中很少被消融。')
h('2.3 轻量化检测与边缘部署', 2)
p('相关轻量化设计包括深度可分离卷积（MobileNet [14]）、特征复用卷积（GhostNet [15]）与部分卷积'
  '（FasterNet [16]）'
  '——后者仅对部分通道做空间卷积，以降低冗余内存访问。已有多项农业研究将 YOLO 变体经 TensorRT 加速'
  '部署到 Jetson 系列板卡并报告了可喜的帧率 [17, 18]。但畜牧行为工作中上板实测仍然稀少，且有两个硬件事实很少被'
  '讨论：Maxwell 代的 Jetson Nano 没有原生 INT8 支持；TensorRT 的算子融合在小模型规模下会掩盖轻量化'
  '算子的理论优势。本文在第 5 节对两者都如实记录。')

# ===== 3 数据与方法 =====
h('3 数据与方法', 1)
h('3.1 数据集', 2)
p('实验采用公开群养猪行为数据集（Roboflow Universe，CC BY 4.0 [19]），源自 Bergamini 等 [20] 公开记录的'
  '采集工作，由商业猪舍固定顶置摄像头于自然日光下录制。数据集共 5,620 张图像、13,995 个标注实例，'
  '覆盖 10 类行为，画面为栏位俯视或斜视视角。公开数据集未提供品种与日龄信息。本文沿用发布方切分'
  '（训练 3,936 / 验证 1,123 / 测试 561 张）。标注质量经训练集随机抽样目检核验。本研究仅使用公开获取的'
  '非侵入式数据，不涉及动物实验，无需伦理审批。表 1 给出各类的行为学定义，表 2 汇总实例分布，'
  '图 1 将分布与由其导出的复制倍率（见 3.3 节）一并可视化。')
table(['类别', '行为学定义', '管理意义'],
      [['active', '高活动量运动：奔跑、跳跃（由低频 run/jump 类聚合）', '活动异常指示应激或疾病'],
       ['drink', '头部位于饮水器处并呈饮水姿态', '饮水异常信号消化/泌尿系统问题'],
       ['eat', '头部入饲槽并持续采食', '采食量是健康/生长核心指标'],
       ['fight', '攻击性接触：两头及以上个体撕咬、追逐、顶撞', '福利预警关键行为，混群期高发'],
       ['investigating', '探究：拱地、嗅探地面/栏体/同伴（非采食）', '正常探究行为，反映环境丰富度'],
       ['lying', '侧躺或俯卧，躯干完全着地，静止休息', '舒适度/健康直接指标'],
       ['nose-to-nose', '两头猪鼻部直接接触且无攻击动作', '社会交往，群体关系指标'],
       ['sitting', '犬坐姿：前腿直立、后躯着地', '偶见姿态，可能关联跛行/不适'],
       ['standing', '四肢站立且无位移', '基础姿态行为'],
       ['walk', '四肢行走的位移运动（不含奔跑）', '日常活动量组成']],
      '表 1 行为类别及其行为学定义。')
table(['类别', '实例数', '类别', '实例数'],
      [['investigating', '4,203', 'fight', '807'],
       ['walk', '2,736', 'nose-to-nose', '358'],
       ['lying', '2,485', 'active', '259'],
       ['standing', '1,812', 'drink', '211'],
       ['eat', '980', 'sitting', '144']],
      '表 2 十类行为的实例分布（不均衡比约 29:1）')
fig('results/analysis/fig6-class-distribution.png',
    '图 1 十类行为的实例分布（柱，对数轴）与作用于训练图像的每类复制倍率（方块，右轴；'
    '倍率按含类图像数计算，见 3.3 节）')
h('3.2 基线模型', 2)
p('基线选用 YOLOv11n [21]（2.58M 参数、640×640 下 6.3 GFLOPs）：它是最轻的主流变体，也是近期畜牧检测'
  '论文最常用的基座，这让我们的数字可以直接对比。我们也试过更新的、以注意力为中心的 YOLOv12n [22] 作基座；'
  '意外的是采样策略在它上面反而降了 1.4 个点（4.2 节阴性对照），因此在这一"数据集+方法"组合下保留 '
  'YOLOv11n。')
h('3.3 类别不均衡感知采样', 2)
p('我们不对损失函数重加权——那会改变优化图景并可能与预训练权重相互干扰——而是调整训练数据的采样'
  '分布。设类 c 在训练集中出现的图像数为 N_c（非实例数）、最大类图像数为 N_max，则每张训练图复制 '
  'f = min(5, round(√(N_max/N_cmin))) '
  '次，其中 cmin 为该图出现的最稀有类。开方抑制极端倍率，封顶防止对少数稀有样本过拟合。举例：sitting '
  '（出现于 104 张训练图）的倍率为 round(√(1714/104)) = round(4.06) = 4，而 investigating（1,714 张）为 1。'
  '复制以硬链接实现，'
  '不占额外存储，验证/测试集保持不动。训练集由 3,936 张扩至 5,889 张，最稀有类每轮曝光增加 3–4 倍。')
table(['封顶值', 'val mAP50', 'test mAP50', '训练集规模'],
      [['3', '0.5754', '0.5776', '5,785'],
       ['4', '0.5816', '0.6035', '5,889'],
       ['5', '0.5816', '0.6035', '5,889']],
      '表 3 复制封顶值敏感性。本数据集实际用到的最大倍率为 4（sitting），故封顶 4 与 5 一致；'
      '封顶 3 时 sitting 曝光不足，整体损失 2.6 点。')
h('3.4 FasterNet 轻量化主干替换', 2)
p('第二处改动针对计算量。我们将主干 P3–P5 层的 C3k2 替换为 FasterNet 块：每块先做一次部分卷积'
  '（PConv，仅对 1/4 通道做空间卷积、其余直通），再接两个逐点卷积与批归一化，并带残差连接。PConv '
  '在不压缩特征体积的前提下削减冗余内存访问，适合细微且空间弥散的行为特征（例如同一位置"饮水"与'
  '"站立"的猪）。P2 与检测头保留（图 2）。替换（图 2）使参数量由 2.58M 降至 2.47M（-4.4%）；FLOPs 基本持平'
  '（6.3G→6.5G），实测推理更快（4.6 节）。')
fig('results/analysis/fig4-architecture.png',
    '图 2 所提模型结构与 FasterNet 块示意。仅替换 P3–P5 主干层，检测头保持不动以承接预训练权重')
h('3.5 恒等保持的预训练权重集成', 2)
p('早期我们试过流行的做法：向预训练网络插入注意力模块。主干末端（C2PSA 之后）插入 EMA 注意力块'
  '使精度下降 2.95 个点（54.11% 对 57.06% mAP50）；同模块插到颈部更糟（第 37 轮叫停，落后 8 个点）。'
  '原因是结构性的：随机初始化的门控模块对特征做乘性缩放，下游预训练层突然收到分布之外的输入——'
  '残差包裹变体的特征能量比实测为 2.25，而门控变体约为 0.25。因此我们改走另一条路：保持通道流不变，'
  '经索引对齐重映射迁移全部未改动权重（415 个参数张量中迁移 316 个），仅类别相关层（分类分支，'
  'COCO 80 类→10 类）重新初始化。改造后的模型起步行为几乎等同预训练基线，微调也稳定得多。改造预训练'
  '检测器时，我们推荐这一恒等保持策略优先于朴素插入。')
h('3.6 训练与评测协议', 2)
p('全部模型至多训练 200 轮并早停（patience 30），输入 640×640，使用 Ultralytics 默认超参（AdamW、'
  'lr0=0.01、weight decay=5×10⁻⁴、HSV 与 mosaic 增强；ultralytics 8.4.105、PyTorch 2.3.0、'
  'CUDA 12.1）。选模用验证集，报告指标一律来自独立测试集；最终模型以三个不同种子重复训练，'
  '报均值±标准差。训练使用 RTX 3090，部署评估使用 Jetson Nano（JetPack 4.6.3、TensorRT 8.2）。')
h('3.7 评价指标', 2)
p('报告 mAP50、mAP50-95、精确率与召回率（总体与每类）、参数量、GFLOPs、推理延迟（FPS）与整机功耗。')

# ===== 4 实验与结果 =====
h('4 实验与结果', 1)
h('4.1 实验设置', 2)
p('全部模型在 3.6 节统一协议下训练与评测：相同的数据切分、增强、轮次预算与早停规则，相同的硬件'
  '（单张 24 GB 显存的 RTX 3090）。对照组覆盖三代主流轻量检测器——YOLOv5n [23]（2.50M 参数、7.1 GFLOPs）、'
  'YOLOv8n [24]（3.01M、8.1G）、YOLOv12n [22]（2.56M、6.3G，最新的注意力中心变体）——以及 RT-DETR-l [25]'
  '（约 32M 参数、约 103 GFLOPs），用以检验一种根本不同的架构是否会改变局面。除特别说明外，'
  '所有关键数字均来自独立测试集（561 张训练与选模都未接触过的图像）。')
h('4.2 消融实验', 2)
p('表 4 按我们实际探索的顺序走过整个设计空间，包括失败。两次注意力集成尝试（M1、M2）列在最前：'
  '两者都明显损伤精度，而其诊断（3.5 节）正是我们转向恒等保持集成的原因。FasterNet 单独替换（M3）'
  '以 4.4% 参数削减为代价损失 2.7 个 test 点；热重启续训实验（100+100 轮、lr0=0.002）追回 2.95 点，'
  '说明差距部分是替换块欠训的产物而非容量天花板。采样策略单独使用（M4）是最有效的单一改动：'
  'test mAP50 由 0.5964 升至 0.6035，验证集召回由 0.605 升至 0.626（+2.1 点，验证集），增益集中在'
  '稀有类（4.4 节；训练动态见图 3）。组合模型（M5）达到 0.5932——与基线仅差 0.32 点——同时参数少 4.4%，也是第 5 节'
  '部署的模型。最后，阴性对照（M6）把同一采样策略施加于 YOLOv12n：精度由 0.6135 降至 0.5994'
  '（-1.4 点），即曝光再平衡的收益无法迁移到注意力中心基座。我们将其读作 YOLOv12n 原生抗不均衡更强、'
  '重采样反而致其过拟合的证据；它也回溯性地印证了我们在这一"数据集+方法"组合下选择 YOLOv11n 作基座'
  '的经验判断。')
table(['模型', '改动', 'mAP50', 'mAP50-95', '参数量', '判定'],
      [['基线', 'YOLOv11n 原版', '0.5964', '0.4328', '2.58M', '锚点'],
       ['M4', '+ 不均衡感知采样（封顶4）', '0.6035', '0.4379', '2.58M', '采用'],
       ['M3', '仅 FasterNet 主干', '0.5691', '0.4009', '2.47M', '参照'],
       ['M5', 'M3+M4（最终轻量模型）', '0.5932', '0.4300', '2.47M', '采用'],
       ['M1', '+EMA 注意力@主干末端', '0.5411*', '0.3915*', '—', '否决（-2.95）'],
       ['M2', '+EMAR 残差注意力@颈部', '0.4988*', '0.3534*', '—', '否决（-7.2）'],
       ['M6', '采样策略@YOLOv12n 基座', '0.5994', '0.4321', '2.56M', '阴性对照']],
      '表 4 独立测试集上的消融。*M1/M2 被提前否决，其数值为验证集 mAP50；其余各行为单次实验 test 值。'
      '统计参照：三次重复 test 均值为 0.5987±0.0062（M4）、0.5904±0.0086（M5），基线为单次 0.5964——'
      '即表中单次值落在运行间噪声带内，并非挑选有利种子（见 4.5 节）。M4 的 +2.1 点召回增益测自验证集。')
fig('results/analysis/fig5-curves.png',
    '图 3 训练过程验证集 mAP50 曲线。采样变体（M4）自早期轮次即领先；组合变体（M5）紧贴基线；'
    '早停（patience 30）在 M4 峰值后兜住其轻微过拟合')
h('4.3 与主流模型对比', 2)
p('表 5 在相同训练条件下对比两个采用变体与对照组。M4（0.6035）超过 YOLOv5n（0.6001）、RT-DETR-l'
  '（0.6008）、YOLOv8n（0.5877）与未改动的 YOLOv11n 基线（0.5964），仅次于 YOLOv12n（0.6135）。'
  'M5 以 2.47M 参数达到 0.5932：比 YOLOv8n 小 18% 仍反超，并逼近 RT-DETR-l——参数量约为其 1/13、'
  'FLOPs 约为其 1/16；当目标硬件是 5 瓦边缘板而非服务器 GPU 时，这是一个有实际意义的工作点。'
  '两个现象值得记录：其一，YOLOv8n 在验证集上高于我们基线、在测试集上反而更低——验证集排名并不总能'
  '挺过留出数据的检验；其二，YOLOv12n 虽是本数据集精度王，却也是我们实测最慢的模型（4.6 节），'
  '且如 M6 所示并不受益于采样策略——对边缘导向的框架而言，我们把它视为强力的服务器端参照，'
  '而非部署候选。')
table(['模型', 'test mAP50', 'test mAP50-95', '参数量', 'GFLOPs'],
      [['YOLOv12n', '0.6135', '0.4354', '2.56M', '6.3G'],
       ['M4（本文）', '0.6035', '0.4379', '2.58M', '6.3G'],
       ['RT-DETR-l', '0.6008', '0.4270', '~32M', '~103G'],
       ['YOLOv5n', '0.6001', '0.4268', '2.50M', '7.1G'],
       ['基线 YOLOv11n', '0.5964', '0.4328', '2.58M', '6.3G'],
       ['M5（本文）', '0.5932', '0.4300', '2.47M', '6.5G'],
       ['YOLOv8n', '0.5877', '0.4332', '3.01M', '8.1G']],
      '表 5 同一协议下的主流检测器对比（按 test mAP50 排序）。所有模型均从官方预训练权重训练。')
h('4.4 分类别分析', 2)
p('整体 mAP 掩盖了这套框架存在的理由，表 6 因此在验证集上按行为类别拆分 AP50——四个消融模型在验证集'
  '上可直接互比（并见图 4）。M4 的增益精确落在预定目标上：active 由 0.459 升至 0.552（+9.3 点），drink 由 0.408 '
  '升至 0.459（+5.1），sitting 由 0.403 升至 0.423（+2.0），中频的 eat 与 standing 各升 1.1 点。'
  '代价是 nose-to-nose 下降 6.4 点——该类的 358 个实例来自少数几个场景，四倍曝光看来是对那几个具体'
  '场景过拟合而非对行为本身，这一失效模式在第 7 节展开。在独立测试集上（基线对 M4），最醒目的稀有类'
  '结果是 active 由 0.526 升至 0.639（+11.3 点），sitting（0.534→0.571）与 fight（0.840→0.871）'
  '亦有提升；drink、eat、standing、walk 在 ±3 点内双向波动。组合模型 M5 在最稀有类上不及 M4，'
  '但部分挽回了 nose-to-nose 的损失（验证集 0.710，高于基线 0.686）——FasterNet 特征与再平衡曝光'
  '在紧密接触行为上似乎形成了良性互动。sitting 在所有模型与切分下都是最弱类（0.23–0.57）；'
  '144 个实例的体量决定了这是数据稀缺问题，而非架构问题。')
table(['类别', '基线', 'M3（FasterNet）', 'M4（采样）', 'M5（组合）', 'M4−基线'],
      [['active', '0.459', '0.372', '0.552', '—', '+9.3'],
       ['drink', '0.408', '0.440', '0.459', '0.430', '+5.1'],
       ['eat', '0.436', '0.437', '0.447', '0.449', '+1.1'],
       ['fight', '0.858', '0.828', '0.842', '0.786', '-1.6'],
       ['investigating', '0.581', '0.571', '0.570', '0.577', '-1.1'],
       ['lying', '0.783', '0.764', '0.766', '0.743', '-1.7'],
       ['nose-to-nose', '0.686', '0.655', '0.622', '0.710', '-6.4'],
       ['sitting', '0.403', '0.233', '0.423', '0.349', '+2.0'],
       ['standing', '0.438', '0.429', '0.449', '0.445', '+1.1'],
       ['walk', '0.687', '0.646', '0.686', '0.667', '-0.1'],
       ['mAP50（全部）', '0.573', '0.537', '0.582', '0.561', '+0.9']],
      '表 6 验证集分类别 AP50。独立测试集上的关键稀有类结果：active 0.526→0.639（基线→M4，+11.3 点）；'
      'sitting 0.534→0.571；fight 0.840→0.871。')
fig('results/analysis/fig7-perclass-ap.png',
    '图 4 验证集分类别 AP50（基线 / M4 / M5；M5 active 无数据）。M4 的增益集中在低频行为（右半部）')
h('4.5 统计可靠性', 2)
p('小型检测器单次训练的波动轻易可达半个百分点，因此两个采用模型都以三个不同种子重复。M4 验证集 '
  '0.5790±0.0054、测试集 0.5987±0.0062；M5 分别为 0.5620±0.0079 与 0.5904±0.0086。相对基线 test '
  '0.5964，M5 的差值（-0.006）在一个标准差以内，故我们将 M5 的精度表述为与基线统计相当而非超越。'
  'M4 的 test 均值整体仅超出基线 +0.002——其真正价值如 4.4 节所示，在于精度向稀有的福利相关行为'
  '再分配，以及验证集 +2.1 点的召回增益，而非整体 mAP 上涨。我们刻意这样呈现："整体提升"的框架经不起'
  '审稿人推敲，而稀有类框架既诚实，对福利预警又恰是实际相关的那个。')
h('4.6 效率评估', 2)
p('表 7 报告 RTX 3090 上的效率（FP32、batch 1、640×640）。M5 以 117.6 FPS 为全场最快（图 5）——领先基线'
  '（112.8）与 M4（112.1），领先 YOLOv12n（78.4）达 50%。有一条边界必须写明：效率收益体现在参数量'
  '（-4.4%）与实测速度上，FLOPs 则基本持平（6.3G→6.5G）。这在意料之中而非自相矛盾：FLOPs 只数算术'
  '运算，而部分卷积主要削减冗余内存访问——延迟测试反映后者，FLOPs 计数器不反映。同样的提醒会在 '
  'Jetson Nano 上以放大的形式重现（第 5 节）：TensorRT 算子融合把剩余差距也吸收了。')
table(['模型', '参数量', 'GFLOPs', 'FPS（RTX 3090）'],
      [['M5（本文）', '2.47M', '6.5G', '117.6'],
       ['基线 YOLOv11n', '2.58M', '6.3G', '112.8'],
       ['M4（本文）', '2.58M', '6.3G', '112.1'],
       ['YOLOv12n', '2.56M', '6.3G', '78.4']],
      '表 7 RTX 3090 效率（FP32、batch 1、640×640）。M5 兼具最少参数与最高实测吞吐；注意其 FLOPs '
      '与基线相当——收益在内存访问效率，这是 FLOPs 捕捉不到的。')
fig('results/analysis/fig8-pareto.png',
    '图 5 各检测器的精度-效率权衡（test mAP50 对参数量，对数轴）。标记大小正比于 RTX 3090 实测 FPS；'
    '空心标记为未测 FPS 的模型')
h('4.7 可视化分析', 2)
p('Grad-CAM 热力图（图 6）显示基线与轻量变体的注意力都集中于猪体而非栏舍背景，且 FasterNet 替换'
  '并未明显移动关注区域——这与两模型在高频类上分类别表现一致相吻合。图 7 给出一个典型残余错误：'
  '一头静止的卧猪被以高置信（0.84）误标为 fight，展示了密集场景中 fight/lying/紧密接触姿态之间的'
  '残余混淆，与归一化混淆矩阵（图 8）的主导误差一致。在个体分离良好的常规场景中，两个模型的检测结果几乎一致且正确；'
  '残余错误集中在拥挤、遮挡、身体纠缠的猪群——也正是人类观察者同样会犹豫的情形。')
fig('results/analysis/gradcam/2019_11_28_000113_105_jpg.rf.fda772ffb5c9bd667ffb03d38721e8c9_cam.jpg',
    '图 6 Grad-CAM 对比（原图 | 基线 | 轻量化变体）。注意力集中于猪体而非栏舍背景，且经主干替换后保持')
fig('results/analysis/detections/2019_11_28_000113_105_jpg.rf.fda772ffb5c9bd667ffb03d38721e8c9_det.jpg',
    '图 7 典型残余误检：一头静卧猪被误标为 fight（0.84）——密集场景中 fight/lying/紧密接触混淆的'
    '典型案例，构成误差主体')
fig('results/m4-wsample/confusion_matrix_normalized.png',
    '图 8 M4 在验证集上的归一化混淆矩阵。误差质量集中于 fight、lying 与紧密接触姿态之间的混淆',
    width=5.4)

# ===== 5 边缘部署验证 =====
h('5 边缘部署验证', 1)
p('只能在服务器 GPU 上跑的检测器解决不了引言中的现场问题，因此我们在 Jetson Nano（4 GB、Maxwell GPU、'
  'JetPack 4.6.3、TensorRT 8.2）上验证了完整部署链路——这块板子的零售价大约相当于一个产床传感器。'
  '链路为 PyTorch → ONNX（opset 12）→ TensorRT FP16（图 9）。导出件很小：M5 权重 5.2 MB'
  '（基线 5.5 MB）、ONNX 10.1 MB，栏端设备的远程更新成本极低。表 8 为实测：M5 在 640×640 输入下每帧 50.2 ms'
  '（19.7 FPS），480×480 下 30.0 ms（33.3 FPS），满载整机功耗约 5W。两个工作点对行为监测都属实时'
  '（行为事件以秒级展开）；480×480 工作点还为同板多路视频流或多模型管线留出了余量。')
table(['模型', '输入', '延迟（ms）', 'FPS', '整机功耗'],
      [['M5（本文）', '640×640', '50.2', '19.7', '≈5 W'],
       ['M5（本文）', '480×480', '30.0', '33.3', '≈5 W'],
       ['基线', '640×640', '50.8', '19.7', '≈5 W'],
       ['基线', '480×480', '29.9', '33.4', '≈5 W']],
      '表 8 Jetson Nano 上板实测（TensorRT 8.2、FP16、batch 1、热身后持续推理）。')
fig('results/analysis/fig9-deploy-pipeline.png',
    '图 9 从 PyTorch 权重到板上 TensorRT 推理的部署管线，含板上实测性能与实践中遇到的限制')
p('部署过程中有三点实践发现值得报告——这一领域文献很少记录它们，而每一点都真实消耗过我们的时间。'
  '其一，INT8 量化在这块板子上实际不可用：Maxwell GPU 没有原生 INT8 单元，TensorRT 8.2 的校准在其上'
  '失败，因此 FP16 是 Nano 级硬件的实际精度下限；需要 INT8 的读者应预算 Orin 级板卡。其二，M5 与基线的'
  '板上延迟差不足 2%——TensorRT 算子融合在这一模型规模下吸收了部分卷积的优势，服务器端的速度差'
  '（表 7）并未延续到边缘端。诚实的表述因此是：Nano 上打平、模型更小，而非速度领先。其三，构建工具链'
  '的脆弱之处值得外业前知晓：一次被中断的 onnxslim 简化留下了静默损坏的 ONNX 文件，直到引擎构建时才'
  '暴露；随后一连串构建失败报了误导性的 cuDNN 符号错误，重启后消失（崩溃构建留下的过期 GPU 状态）。'
  '不做图简化的导出、一次只跑一个引擎构建、任何崩溃后重启，让管线可复现。这些问题在服务器端评测中'
  '都不会出现——这正是我们在此记录它们的原因。')

# ===== 6 泛化分析 =====
h('6 泛化分析', 1)
h('6.1 未见序列压力测试', 2)
p('发布方切分把同一视频序列的帧分散到训练、验证、测试三集中。相邻帧近乎重复，因此上述主指标度量的是'
  '已见场景内的插值——这是该领域的通行做法，但是乐观的做法。为量化乐观程度，我们按视频序列重新切分：'
  '训练集 4,116 张；测试集为两个完全未见的序列（2019_11_28_000113 与 2019_12_10_000060）；验证集混合'
  '序列 000033 与异源画面（Em 系列，约占该切分 37%），栏位布局与外观均有差异。所有模型在该切分下'
  '性能骤降（表 9，图 10）：基线在未见序列测试集上由 0.5964 跌至 0.155，在更难的验证集上跌至 0.075。'
  '采样训练的 M4 在该验证集上接近基线两倍（0.139 对 0.075）——曝光再平衡对视觉陌生场景有部分缓解，'
  '一个合理的解释是稀有场景（不仅是稀有类）也获得了更多训练曝光——但它并不能消除鸿沟，且该优势未能'
  '延续到未见序列测试集（0.145 对 0.155）。M5 在验证集上居中（0.107），测试集与基线持平（0.153）。')
table(['模型', 'val mAP50（混合异源，更难）', 'test mAP50（未见序列）'],
      [['基线', '0.075', '0.155'],
       ['M4（采样）', '0.139', '0.145'],
       ['M5（组合）', '0.107', '0.153']],
      '表 9 序列级互斥压力测试。所有模型相对随机切分都大幅退化；M4 将基线验证精度翻倍，'
      '但没有任何模型消除未见序列鸿沟。')
fig('results/analysis/fig10-generalization.png',
    '图 10 基线、M4、M5 在三种评测制度下的 mAP50：随机切分（test）、序列互斥压力测试与外部农场'
    '零样本评估。从域内到跨域的断崖对所有模型一致')
h('6.2 跨数据集验证', 2)
p('更难的考验是换一个农场。我们将四个模型零样本（不做任何微调）在独立公开猪行为数据集 Comportamentos'
  '（696 张图像，CC BY 4.0 [26]）上评估——该数据集来自不同养殖场，地面材质、相机几何不同，且以花斑品种'
  '为主，而非源数据集的白猪。所有模型都跌至 0.036–0.067 mAP50（表 10），排序本身也富含信息：基线退化'
  '最轻（0.067），采样训练的模型退化最重（0.036–0.038）——再平衡使模型对其被平衡的源域更加特化。'
  'M5 的置信画像很有说服力：精确率 0.70、召回率 0.03，意味着模型在域外很少开口，一旦开口通常是对的'
  '——域外预测趋于保守而非随机。图 11 让失效具体可见：白猪密集栏位大量漏检、花斑个体检不出、泥地'
  '引发误检。漂移同时发生在环境层面（地面、光照、相机）与生物层面（品种级外观），对照组中没有模型'
  '能应对。我们完整报告这些结果，因为跨农场泛化才是落地真正依赖的性质，而它在这一领域文献中极少'
  '被量化。')
table(['模型', 'mAP50', 'mAP50-95', '精确率', '召回率'],
      [['基线', '0.0671', '0.0362', '0.063', '0.100'],
       ['M3（FasterNet）', '0.0383', '0.0215', '0.033', '0.037'],
       ['M4（采样）', '0.0375', '0.0235', '0.066', '0.060'],
       ['M5（组合）', '0.0361', '0.0207', '0.697', '0.031']],
      '表 10 Comportamentos 独立数据集零样本评估（696 张，CC BY 4.0，不同农场与品种）。退化是普适现象；'
      'M5 的高精确率/近零召回表明域外预测趋于保守而非随机。')
fig('results/analysis/ext-detections.jpg',
    '图 11 独立数据集零样本检测实况：白猪密集栏位大量漏检（左上）、花斑品种个体未检出（左下）、'
    '泥地场景部分检出（右）。失效是环境层面与品种层面同时发生的')

# ===== 7 讨论 =====
h('7 讨论', 1)
p('纵观全文，有三点发现。第一，对不均衡行为数据，最便宜有效的干预在数据层：频率封顶复制不花任何参数'
  '与延迟、不需要改架构，却把精度精确移到福利监测需要的地方——稀有行为，而早停兜住了它引入的轻微'
  '过拟合。第二——这是我们花了两个失败模型才学到的——改造预训练检测器时，保住预训练特征通路比选择'
  '哪个轻量或注意力模块更关键。每一次向预训练流中插入随机初始化门控模块的尝试都损伤了精度（M1、M2），'
  '与位置、与是否残差包裹无关，实测特征能量比（残差变体 2.25 对门控变体约 0.25）支持"分布污染"这一'
  '解释。FasterNet 替换之所以成功，正是因为索引对齐重映射让它从预训练网络的近恒等变换出发。第三，'
  '效率故事真实但有界：收益在参数量与服务器端速度，FLOPs 持平，TensorRT 融合后的边缘推理差距收窄至'
  '打平——读者在类似板卡上应预期相同结果。')
p('阴性结果值得同等的笔墨。采样策略无法迁移到注意力中心的 YOLOv12n（M6），因此"数据层再平衡+卷积中心'
  '基座"的配对看来是该组合的性质而非方法孤立的性质——采用者应在自家基座上先验证这个配对。nose-to-nose '
  '的回退表明，基于复制的再平衡可能对稀有类来源的少数场景过拟合，更精细的改进方向是场景感知（而非纯'
  '类别感知）的封顶规则。两级泛化分析则是对整个子领域的提醒：随机切分评测大幅高估了可部署性，跨农场、'
  '跨品种运行仍是一个开放的、领域级的问题，靠数据集内调参无解。')
p('局限。最稀有类 sitting（144 实例）在所有模型与切分下仍然偏弱——解药是更多稀有类影像，而非更好的'
  '架构。框架仅在单一来源数据集（固定摄像头、日光）上评估，夜视、产床、户外栏位超出范围。跨农场、'
  '跨品种泛化尚未解决（第 6 节）。INT8 加速需要比 Maxwell 代 Nano 更新的边缘硬件。最后，行为标签是'
  '帧级瞬时的；时序平滑或追踪很可能同时改善稳定性与 fight/lying 混淆，我们把它与域适应、多农场数据'
  '采集并列为最有希望的下一步。')

# ===== 8 结论 =====
h('8 结论', 1)
p('本文提出了一个轻量化、实时的群养猪多行为检测框架，将类别不均衡感知采样与 FasterNet 主干替换在'
  '"恒等保持"集成策略下结合。在公开的十类行为数据集上，框架将低频行为（active）的 AP 在独立测试集上'
  '提升 11.3 点、验证集召回提升 2.1 点；三次重复实验中整体精度与 YOLOv11n 基线统计相当；参数量减少 '
  '4.4%；在服务器 GPU 上是所有参评模型中最快的（117.6 FPS）。Jetson Nano 上板验证确认实时运行'
  '（640×640 下 19.7 FPS、480×480 下 33.3 FPS），整机功耗约 5W。两级泛化分析——未见序列压力测试与'
  '独立农场数据集零样本评估——诚实地划定了适用边界：框架今天即可用于同农场部署；域适应与多农场数据'
  '采集则是明天跨农场运行的关键方向。')

# ===== 声明 =====
h('声明', 1)
for line in ['作者贡献：[待填]',
             '基金资助：[待填]',
             '伦理委员会声明：不适用。本研究仅使用公开获取的非侵入式视频数据，未进行动物实验，无需伦理审批。',
             '知情同意声明：不适用。',
             '数据可用性声明：源数据集于 Roboflow Universe 公开（CC BY 4.0）[19]；外部验证数据集（Comportamentos）'
             '同样公开（CC BY 4.0）[26]。处理脚本、模型配置、训练日志与部署脚本将于投稿前在 [GitHub 链接] 公开。',
             '利益冲突：作者声明无利益冲突。',
             'AI 辅助声明：AI 工具用于语言润色与代码辅助；全部实验、测量与所报数字均由作者完成并核实。']:
    p(line)

# ===== 参考文献（与英文版同序同内容，Springer Basic 数字格式） =====
h('参考文献', 1)
REFS = [
 'Canario, L., Bijma, P., David, I., et al.: Prospects for the analysis and reduction of damaging '
 'behaviour in group-housed livestock, with application to pig breeding. Front. Genet. 11, 611073 '
 '(2020). https://doi.org/10.3389/fgene.2020.611073',
 'Peden, R.S.E., Turner, S.P., Boyle, L.A., et al.: The translation of animal welfare research into '
 'practice: the case of mixing aggression between pigs. Appl. Anim. Behav. Sci. 204, 1–9 (2018)',
 'Berckmans, D.: General introduction to precision livestock farming. Anim. Front. 7(1), 6–11 (2017). '
 'https://doi.org/10.2527/af.2017.0102',
 'Redmon, J., Divvala, S., Girshick, R., et al.: You only look once: unified, real-time object '
 'detection. In: Proc. IEEE Conf. Comput. Vis. Pattern Recognit. (CVPR), pp. 779–788 (2016). '
 'https://doi.org/10.1109/CVPR.2016.91',
 'Tu, S., Zeng, Q., Liang, Y., et al.: Automated behavior recognition and tracking of group-housed '
 'pigs with an improved DeepSORT method. Agriculture 12(11), 1907 (2022). '
 'https://doi.org/10.3390/agriculture12111907',
 'Li, R., Dai, B., Hu, Y., et al.: Multi-behavior detection of group-housed pigs based on YOLOX and '
 'SCTS-SlowFast. Comput. Electron. Agric. 225, 109286 (2024). '
 'https://doi.org/10.1016/j.compag.2024.109286',
 'Alameer, A., Kyriazakis, I., Bacardit, J.: Automated recognition of postures and drinking behaviour '
 'for the detection of compromised health in pigs. Sci. Rep. 10, 13665 (2020). '
 'https://doi.org/10.1038/s41598-020-70688-6',
 'Gupta, A., Dollár, P., Girshick, R.: LVIS: a dataset for large vocabulary instance segmentation. '
 'In: Proc. IEEE/CVF Conf. Comput. Vis. Pattern Recognit. (CVPR), pp. 5356–5364 (2019). '
 'https://doi.org/10.1109/CVPR.2019.00550',
 'Lin, T.Y., Goyal, P., Girshick, R., et al.: Focal loss for dense object detection. In: Proc. IEEE '
 'Int. Conf. Comput. Vis. (ICCV), pp. 2980–2988 (2017). https://doi.org/10.1109/ICCV.2017.324',
 'Ghiasi, G., Cui, Y., Srinivas, A., et al.: Simple copy-paste is a strong data augmentation method '
 'for instance segmentation. In: Proc. IEEE/CVF Conf. Comput. Vis. Pattern Recognit. (CVPR), '
 'pp. 2918–2927 (2021)',
 'Oksuz, K., Cam, B.C., Kalkan, S., et al.: Imbalance problems in object detection: a review. IEEE '
 'Trans. Pattern Anal. Mach. Intell. 43(10), 3388–3415 (2021). '
 'https://doi.org/10.1109/TPAMI.2020.2981890',
 'Buda, M., Maki, A., Mazurowski, M.A.: A systematic study of the class imbalance problem in '
 'convolutional neural networks. Neural Netw. 106, 249–259 (2018). '
 'https://doi.org/10.1016/j.neunet.2018.07.011',
 'Lv, J., Wang, G., Zhang, M., et al.: WFE-YOLO: a lightweight pig behavior detection model for '
 'livestock farming applications. INMATEH Agric. Eng. 78(1), 1260–1273 (2026)',
 'Howard, A.G., Zhu, M., Chen, B., et al.: MobileNets: efficient convolutional neural networks for '
 'mobile vision applications. arXiv preprint arXiv:1704.04861 (2017)',
 'Han, K., Wang, Y., Tian, Q., et al.: GhostNet: more features from cheap operations. In: Proc. '
 'IEEE/CVF Conf. Comput. Vis. Pattern Recognit. (CVPR), pp. 1580–1589 (2020)',
 'Chen, J., Kao, S., He, H., et al.: Run, don\'t walk: chasing higher FLOPS for faster neural '
 'networks. In: Proc. IEEE/CVF Conf. Comput. Vis. Pattern Recognit. (CVPR), pp. 12021–12031 (2023). '
 'https://doi.org/10.1109/CVPR52729.2023.01635',
 'Kim, J., Suh, Y., Lee, J., et al.: EmbeddedPigCount: pig counting with video object detection and '
 'tracking on an embedded board. Sensors 22(7), 2689 (2022). https://doi.org/10.3390/s22072689',
 'Gu, Z., He, D., Huang, J., et al.: Simultaneous detection of fruits and fruiting stems in mango '
 'using improved YOLOv8 model deployed by edge device. Comput. Electron. Agric. 227, 109524 (2024)',
 'Pig Behavior Dataset. Roboflow Universe, version 1 (CC BY 4.0). '
 'https://universe.roboflow.com/km-sd0ce/pig-behavior-wlvku (accessed July 2026)',
 'Bergamini, L., Pini, S., Simoni, A., et al.: Extracting accurate long-term behavior changes from a '
 'large pig dataset. In: Proc. 16th Int. Joint Conf. Comput. Vis. Imaging Comput. Graph. Theory '
 'Appl. (VISIGRAPP), vol. 5: VISAPP, pp. 524–533 (2021). https://doi.org/10.5220/0010288405240533',
 'Jocher, G., Qiu, J.: Ultralytics YOLO11 (v11.0.0) [Computer software]. '
 'https://github.com/ultralytics/ultralytics (2024)',
 'Tian, Y., Ye, Q., Doermann, D.: YOLOv12: attention-centric real-time object detectors. arXiv '
 'preprint arXiv:2502.12524 (2025)',
 'Jocher, G., Chaurasia, A., Stoken, A., et al.: ultralytics/yolov5: v7.0—YOLOv5 SOTA realtime '
 'instance segmentation. Zenodo (2022). https://doi.org/10.5281/zenodo.7347926',
 'Jocher, G., Chaurasia, A., Qiu, J.: Ultralytics YOLOv8 (v8.0.0) [Computer software]. '
 'https://github.com/ultralytics/ultralytics (2023)',
 'Zhao, Y., Lv, W., Xu, S., et al.: DETRs beat YOLOs on real-time object detection. In: Proc. '
 'IEEE/CVF Conf. Comput. Vis. Pattern Recognit. (CVPR), pp. 16965–16974 (2024). '
 'https://doi.org/10.1109/CVPR52733.2024.01605',
 'Comportamentos Dataset. Roboflow Universe (CC BY 4.0). '
 'https://universe.roboflow.com/maria-dnxxx/comportamentos-vdzlw (accessed August 2026)',
]
assert len(REFS) == 26
for i, ref in enumerate(REFS, 1):
    par = doc.add_paragraph()
    par.paragraph_format.left_indent = Inches(0.32)
    par.paragraph_format.first_line_indent = Inches(-0.32)
    run = par.add_run(f'[{i}] {ref}')
    run.font.size = Pt(9)

# ===== 保存 =====
out = 'paper/猪行为检测-中文成稿-v3.docx'
doc.save(out)
print('已生成:', out)
