# v5 小修（报告#8 R2-Minor 两项）：3.1 标注协议出处一句 + 讨论章告警路径一句（中英双版锚点）
# 运行：python scripts/apply_v5_r2minor.py  → 原地更新 paper/JRTIP-paper-v5.docx
import sys
sys.stdout.reconfigure(encoding='utf-8', errors='replace')
from docx import Document

PATH = 'paper/JRTIP-paper-v5.docx'
doc = Document(PATH)
paras = doc.paragraphs

def swap(idx_pred, old, new):
    i = next(i for i, p in enumerate(paras) if idx_pred(p.text))
    p = paras[i]
    txt = p.text
    assert old in txt, f'锚点未命中: {old[:50]}'
    nt = txt.replace(old, new)
    for r in list(p.runs):
        r.text = ''
    (p.runs[0] if p.runs else p.add_run()).text = nt
    print(f'OK para {i}: +{new[-60:]}')

# ① 3.1 标注协议出处（挂在标注质量句后）
swap(lambda t: t.startswith('We use a public group-housed pig behavior dataset'),
     'Annotation quality was verified by visual inspection of a random sample of the training images.',
     'Behavior definitions and the annotation protocol follow the original acquisition study [20]; '
     'annotation quality was additionally verified by visual inspection of a random sample of the '
     'training images.')

# ② 讨论章告警路径（Limitations 段，时序平滑句之后）
swap(lambda t: t.startswith('Limitations.'),
     'Temporal integration remains promising for dense-frame live video—precisely the deployment '
     'regime of Section 5—together with domain adaptation and multi-farm data collection.',
     'On the alert side, per-frame labels aggregate naturally into pen-level statistics—fight '
     'frequency, lying-time share, drink-visit counts—over which simple threshold rules can raise '
     'welfare alerts; tuning those thresholds is farm-specific and belongs to deployment work. '
     'Temporal integration remains promising for dense-frame live video—precisely the deployment '
     'regime of Section 5—together with domain adaptation and multi-farm data collection.')

doc.save(PATH)

# 复检
full = '\n'.join(p.text for p in Document(PATH).paragraphs)
for probe in ['annotation protocol follow the original acquisition study [20]',
              'pen-level statistics—fight frequency', 'tuning those thresholds is farm-specific']:
    assert probe in full, f'复检失败: {probe[:50]}'
print('复检通过：两处小修落位')
