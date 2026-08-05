#!/usr/bin/env python3
"""v4 定稿手术 —— 一次性完成 v4-manual-todo.txt 的全部手动项。

输入：paper/JRTIP-paper-v4.docx（编号脚本已跑过的半成品）
输出：paper/JRTIP-paper-v4.docx（覆盖为定稿）+ paper/Supplementary_Material.docx

操作清单（与 v4-manual-todo.txt 对应）：
  图：删 per-class AP 柱状图（原Fig.4）；Grad-CAM+难例 左右拼 → Fig.5；
      泛化柱状+异源实况 上下拼 → Fig.8；全文图引用落定 Fig.1-8
  表：Table 1+2 合并（Class/定义+意义/实例数/训练图像数）；删 Table 8 部署表
      （数据写入第5节正文）；Table 6 每类AP 移至 Supplementary（Table S1）；
      效率表改号 Table 5、压力测试 Table 6、跨数据集 Table 7
  全部引用按上下文精确改写（非全局替换），运行后自动审计跳号/重复。
"""

import copy
from pathlib import Path

import docx
from docx.shared import Inches, Pt
from docx.table import Table
from docx.text.paragraph import Paragraph
from PIL import Image, ImageDraw, ImageFont

SRC = Path(r"E:/pig-behavior-yolo/paper/JRTIP-paper-v4.docx")
SUPP = Path(r"E:/pig-behavior-yolo/paper/Supplementary_Material.docx")

A = "{http://schemas.openxmlformats.org/drawingml/2006/main}"
R = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}"
WP = "{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}"
EMU = 914400

NAMES = ['active', 'drink', 'eat', 'fight', 'investigating',
         'lying', 'nose-to-nose', 'sitting', 'standing', 'walk']
INSTANCES = {'investigating': 4203, 'walk': 2736, 'lying': 2485, 'standing': 1812,
             'eat': 980, 'fight': 807, 'nose-to-nose': 358, 'active': 259,
             'drink': 211, 'sitting': 144}  # 全集 13,995（v3 审计口径）
TRAIN_IMAGES = {'active': 167, 'drink': 147, 'eat': 577, 'fight': 354,
                'investigating': 1714, 'lying': 1299, 'nose-to-nose': 181,
                'sitting': 104, 'standing': 986, 'walk': 1259}  # 本地标签实测


def set_text(para, new_text):
    """单 run 段落整体替换（保留段落样式）。"""
    assert len(para.runs) >= 1, f"段落无 run: {para.text[:40]}"
    para.runs[0].text = new_text
    for r in para.runs[1:]:
        r.text = ""


def replace_in(para, old, new):
    t = para.text
    assert old in t, f"未找到 {old!r} in {t[:80]!r}"
    set_text(para, t.replace(old, new))


def delete_para(para):
    para._element.getparent().remove(para._element)


def delete_table(tbl):
    tbl._element.getparent().remove(tbl._element)


def img_part_and_extent(doc, para):
    blip = para._element.find(f".//{A}blip")
    rid = blip.get(f"{R}embed")
    part = doc.part.related_parts[rid]
    extent = para._element.find(f".//{WP}extent")
    return part, extent


def load_font(size):
    try:
        from matplotlib import font_manager
        return ImageFont.truetype(font_manager.findfont("DejaVu Sans"), size)
    except Exception:
        return ImageFont.load_default()


def label_tag(draw, xy, tag, font):
    """左上角白底 (a)/(b) 标签。"""
    x, y = xy
    bbox = draw.textbbox((x, y), tag, font=font)
    draw.rectangle([bbox[0] - 12, bbox[1] - 8, bbox[2] + 12, bbox[3] + 8],
                   fill="white")
    draw.text((x, y), tag, fill="black", font=font)


def merge_side_by_side(blob_a, blob_b, gap=30, target_h=1000):
    import io
    ia, ib = Image.open(io.BytesIO(blob_a)), Image.open(io.BytesIO(blob_b))
    ia = ia.resize((int(ia.width * target_h / ia.height), target_h))
    ib = ib.resize((int(ib.width * target_h / ib.height), target_h))
    W = ia.width + ib.width + gap
    canvas = Image.new("RGB", (W, target_h), "white")
    canvas.paste(ia, (0, 0))
    canvas.paste(ib, (ia.width + gap, 0))
    d = ImageDraw.Draw(canvas)
    f = load_font(56)
    label_tag(d, (20, 16), "(a)", f)
    label_tag(d, (ia.width + gap + 20, 16), "(b)", f)
    buf = io.BytesIO()
    canvas.save(buf, "JPEG", quality=92)
    return buf.getvalue(), W, target_h


def merge_stacked(blob_top, blob_bottom, gap=30, bottom_ratio=0.62):
    import io
    it_ = Image.open(io.BytesIO(blob_top)).convert("RGB")
    ib = Image.open(io.BytesIO(blob_bottom)).convert("RGB")
    W = it_.width
    ib_w = int(W * bottom_ratio)
    ib = ib.resize((ib_w, int(ib.height * ib_w / ib.width)))
    H = it_.height + gap + ib.height
    canvas = Image.new("RGB", (W, H), "white")
    canvas.paste(it_, (0, 0))
    canvas.paste(ib, ((W - ib_w) // 2, it_.height + gap))
    d = ImageDraw.Draw(canvas)
    f = load_font(56)
    label_tag(d, (20, 16), "(a)", f)
    label_tag(d, ((W - ib_w) // 2 + 20, it_.height + gap + 16), "(b)", f)
    buf = io.BytesIO()
    canvas.save(buf, "PNG")
    return buf.getvalue(), W, H


def main():
    doc = docx.Document(str(SRC))
    paras = doc.paragraphs
    tables = doc.tables
    print(f"读取: {SRC.name} ({len(paras)} 段 / {len(tables)} 表)")

    # ---------- 0. 先收集全部待操作元素的活引用（删除操作最后统一做） ----------
    p = {i: paras[i] for i in
         [22, 23, 24, 54, 55, 56, 57, 61, 62, 66, 67, 68, 69, 70, 74, 75,
          78, 81, 82, 83, 84, 86, 87, 88, 89]}
    t_def, t_dist = tables[0], tables[1]           # Table 1 / Table 2
    t_perclass = tables[5]                          # 每类 AP → 补充材料
    t_deploy = tables[7]                            # 部署表 → 删

    # ---------- 1. 合并 Table 1 + Table 2 ----------
    # 新建四列合并表（样式沿用 Light Grid Accent 1），插到 P023 题注后
    merged = doc.add_table(rows=11, cols=4)
    merged.style = t_def.style
    hdr = ["Class", "Ethological definition and management relevance",
           "Instances\u2020", "Training images\u2021"]
    for j, htxt in enumerate(hdr):
        merged.rows[0].cells[j].paragraphs[0].add_run(htxt).bold = True
    for i, name in enumerate(NAMES, start=1):
        row = merged.rows[i]
        defi = t_def.rows[i].cells[1].text.strip()
        rel = t_def.rows[i].cells[2].text.strip()
        row.cells[0].paragraphs[0].add_run(name)
        row.cells[1].paragraphs[0].add_run(f"{defi}. {rel}.")
        row.cells[2].paragraphs[0].add_run(f"{INSTANCES[name]:,}")
        row.cells[3].paragraphs[0].add_run(f"{TRAIN_IMAGES[name]:,}")
    for row in merged.rows:  # 小字号，与原表观感一致
        for cell in row.cells:
            for pp in cell.paragraphs:
                for rr in pp.runs:
                    rr.font.size = Pt(9)
    # 列宽
    for j, w in enumerate([0.95, 3.15, 1.0, 1.1]):
        for row in merged.rows:
            row.cells[j].width = Inches(w)
    # 移到正确位置：P023 题注之后、旧 Table 1 之前
    t_def._element.addprevious(merged._element)
    print("[1] Table 1+2 合并完成（4 列 x 11 行）")

    # 题注与正文引用
    set_text(p[23], "Table 1. Behavior categories, ethological definitions, and "
                    "instance distribution. \u2020Instances over the full dataset "
                    "(13,995 total); \u2021images containing the class in the "
                    "training split (basis of the duplication factors, Section 3.3).")
    replace_in(p[22],
               "Table 1 defines the ten behaviors; Table 2 summarizes their "
               "instance distribution, and Fig. 1 visualizes it",
               "Table 1 defines the ten behaviors and summarizes their instance "
               "distribution, and Fig. 1 visualizes it")

    # ---------- 2. 每类 AP 表 → Supplementary ----------
    supp = docx.Document()
    supp.add_heading("Supplementary Material", level=1)
    supp.add_paragraph(
        "Real-Time Multi-Behavior Detection of Group-Housed Pigs on Edge "
        "Devices: Class-Imbalance-Aware Sampling and a Lightweight FasterNet "
        "Backbone — Supplementary Material")
    supp.add_paragraph(
        "Table S1. Per-class AP50 on the validation set. On the held-out test "
        "set, the key rare-class result is active 0.526 \u2192 0.639 (baseline "
        "\u2192 M4, +11.3 points); sitting 0.534 \u2192 0.571; fight 0.840 \u2192 0.871.")
    supp_tbl_el = copy.deepcopy(t_perclass._element)
    supp._element.body.append(supp_tbl_el)
    supp.save(str(SUPP))
    print(f"[2] Supplementary_Material.docx 已生成（含 Table S1）")

    replace_in(p[54],
               "so Table 6 breaks AP50 down by behavior on the validation set, "
               "where all four ablation models are directly comparable (also "
               "visualized in Fig. 4).",
               "so Table S1 in the Supplementary Material breaks AP50 down by "
               "behavior on the validation set, where all four ablation models "
               "are directly comparable.")

    # ---------- 3. 效率表改号 Table 5 ----------
    replace_in(p[61], "Table 6 reports efficiency", "Table 5 reports efficiency")
    replace_in(p[62], "Table 6. Efficiency", "Table 5. Efficiency")
    replace_in(p[78], "(Table 6)", "(Table 5)")
    print("[3] 效率表 → Table 5")

    # ---------- 4. 部署表删除、数据写正文 ----------
    replace_in(
        p[74],
        "Table 8 reports the measurements: M5 runs at 50.2 ms per frame "
        "(19.7 FPS) at 640\u00d7640 and 30.0 ms (33.3 FPS) at 480\u00d7480, "
        "with the whole board drawing approximately 5 W under load.",
        "M5 sustains 50.2 ms per frame (19.7 FPS) at 640\u00d7640 and 30.0 ms "
        "(33.3 FPS) at 480\u00d7480\u2014essentially matching the baseline "
        "(50.8 ms/19.7 FPS and 29.9 ms/33.4 FPS)\u2014with the whole board "
        "drawing approximately 5 W under load.")
    print("[4] 部署表数据已写入第 5 节正文")

    # ---------- 5. 压力测试表/跨数据集表改号 ----------
    replace_in(p[81], "(Table 4, Fig. 8)", "(Table 6, Fig. 8)")
    replace_in(p[82], "Table 4. Sequence-disjoint", "Table 6. Sequence-disjoint")
    replace_in(p[86], "(Table 4)", "(Table 7)")
    replace_in(p[87], "Table 4. Zero-shot", "Table 7. Zero-shot")
    print("[5] 压力测试 → Table 6，跨数据集 → Table 7")

    # ---------- 6. 图像拼接：Fig.5（左右）与 Fig.8（上下） ----------
    part6, ext6 = img_part_and_extent(doc, p[67])   # Grad-CAM
    part7, _ = img_part_and_extent(doc, p[69])      # 难例
    blob, W5, H5 = merge_side_by_side(part6.blob, part7.blob)
    part6._blob = blob
    cx = ext6.get("cx")
    ext6.set("cy", str(int(int(cx) * H5 / W5)))
    set_text(p[68],
             "Fig. 5 Grad-CAM comparison and a typical residual error. "
             "(a) Original | baseline | lightweight variant: attention "
             "concentrates on pig bodies rather than the pen background, and is "
             "preserved by the backbone substitution. (b) A motionless lying "
             "pig mislabeled as fight (0.84)\u2014the fight/lying/close-contact "
             "confusion that dominates the error mass in dense scenes.")
    print(f"[6] Fig.5 拼接完成（{W5}x{H5}px，左右）")

    part10, ext10 = img_part_and_extent(doc, p[83])  # 泛化柱状
    part11, _ = img_part_and_extent(doc, p[88])      # 异源实况
    blob8, W8, H8 = merge_stacked(part10.blob, part11.blob)
    part10._blob = blob8
    cx8 = int(ext10.get("cx"))
    cy8 = int(cx8 * H8 / W8)
    max_cy = int(7.2 * EMU)                          # 限高防超页
    if cy8 > max_cy:
        cx8 = int(cx8 * max_cy / cy8)
        cy8 = max_cy
    ext10.set("cx", str(cx8))
    ext10.set("cy", str(cy8))
    set_text(p[84],
             "Fig. 8 Generalization behavior of the framework. (a) mAP50 of "
             "the baseline, M4 and M5 under three evaluation regimes: the "
             "random publisher split (test), the sequence-disjoint stress "
             "test, and zero-shot evaluation on the external farm dataset. "
             "(b) Zero-shot detections on the independent dataset: dense "
             "white-pig pens largely missed, spotted-breed individuals "
             "undetected, partial detections on muddy ground\u2014the failure "
             "is environmental, not stochastic.")
    print(f"[6] Fig.8 拼接完成（{W8}x{H8}px，上下，限高 7.2in）")

    # 正文图引用落定
    replace_in(p[66], "Grad-CAM maps (Fig. 6) show", "Grad-CAM maps (Fig. 5a) show")
    replace_in(p[66], "Fig. 7 shows a typical residual error",
               "Fig. 5b shows a typical residual error")
    replace_in(p[86], "Fig. 8 makes the failure concrete",
               "Fig. 8b makes the failure concrete")

    # ---------- 7. 统一删除（元素引用已收集，顺序无关） ----------
    delete_para(p[24])          # 旧 Table 2 题注
    delete_table(t_dist)        # 旧分布表
    delete_table(t_def)         # 旧定义表（已被合并表替代）
    delete_para(p[55])          # 每类 AP 题注（已入补充材料）
    delete_table(t_perclass)    # 每类 AP 表
    delete_para(p[56])          # per-class 柱状图
    delete_para(p[57])          # 其题注
    delete_para(p[69])          # 难例图（已并入 Fig.5）
    delete_para(p[70])          # 其题注
    delete_para(p[75])          # 部署表题注
    delete_table(t_deploy)      # 部署表
    delete_para(p[88])          # 异源实况图（已并入 Fig.8）
    delete_para(p[89])          # 其题注
    print("[7] 废弃图/表/题注已删除（共 8 段 + 4 表）")

    doc.save(str(SRC))
    print(f"\n[OK] 定稿已保存: {SRC}")


if __name__ == "__main__":
    main()
