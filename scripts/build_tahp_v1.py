# 由 JRTIP-paper-v5.docx 生成 TAHP 投稿版（Tropical Animal Health and Production）
# 变换：①正文 [n] 引用 → 作者-年份制 ②文献表 Springer 作者-年份格式 + 字母序
#       ③结论并入讨论 ④Statement of Animal Rights ⑤A4+双倍行距+连续行号
# 运行：python scripts/build_tahp_v1.py → paper/TAHP-paper-v1.docx
import re
import sys
sys.stdout.reconfigure(encoding='utf-8', errors='replace')
from docx import Document
from docx.shared import Pt, Mm

SRC = 'paper/JRTIP-paper-v5.docx'
DST = 'paper/TAHP-paper-v1.docx'

# [n] → (文内作者, 年份)；Roboflow 两个数据集用 2026a/b；Luo 两篇 2025a/b
REFKEY = {
    1: ('Canario et al.', 2020), 2: ('Peden et al.', 2018), 3: ('Berckmans', 2017),
    4: ('Redmon et al.', 2016), 5: ('Tu et al.', 2022), 6: ('Li et al.', 2024),
    7: ('Alameer et al.', 2020), 8: ('Gupta et al.', 2019), 9: ('Lin et al.', 2017),
    10: ('Ghiasi et al.', 2021), 11: ('Oksuz et al.', 2021), 12: ('Buda et al.', 2018),
    13: ('Lv et al.', 2026), 14: ('Howard et al.', 2017), 15: ('Han et al.', 2020),
    16: ('Chen et al.', 2023), 17: ('Kim et al.', 2022), 18: ('Gu et al.', 2024),
    19: ('Roboflow', '2026a'), 20: ('Bergamini et al.', 2021), 21: ('Jocher and Qiu', 2024),
    22: ('Tian et al.', 2025), 23: ('Jocher et al.', 2022), 24: ('Jocher et al.', 2023),
    25: ('Zhao et al.', 2024), 26: ('Roboflow', '2026b'), 27: ('Luo et al.', '2025b'),
    28: ('Rahman et al.', 2026), 29: ('Yaman et al.', 2023), 30: ('Crasto', 2024),
    31: ('Guo et al.', 2024), 32: ('Qin and Zhou', 2024), 33: ('Luo et al.', '2025a'),
}


def author_year(n):
    a, y = REFKEY[n]
    return f'{a} {y}'


doc = Document(SRC)
paras = doc.paragraphs
i_ref = next(i for i, p in enumerate(paras) if p.text.strip() == 'References')

# ---------- ① 正文引用转换（引用区之前的段落） ----------
CITE_RE = re.compile(r'\[([\d,\s–-]+)\]')


def conv_text(t):
    # A. "Name et al. [n]" / "ModelName [n]" → "Name et al. (year)" / "ModelName (Author year)"
    def repl_named(m):
        head, grp = m.group(1), m.group(2)
        nums = [int(x) for x in re.findall(r'\d+', grp)]
        # 作者已具名（et al. 结尾）：只补年份
        if head.endswith('et al.'):
            years = '; '.join(str(REFKEY[n][1]) for n in nums)
            return f'{head} ({years})'
        # 模型名/简称结尾：补全 作者+年份
        full = '; '.join(author_year(n) for n in nums)
        return f'{head} ({full})'
    t = CITE_RE.sub(lambda m: None, t) if False else t  # 占位（防误用）
    t = re.sub(r'([A-Za-z0-9\-]+(?: et al\.)?)\s+\[([\d,\s–-]+)\]', repl_named, t)
    # B. 剩余独立 [n] / [n, m] → (Author year; Author year)
    def repl_standalone(m):
        nums = [int(x) for x in re.findall(r'\d+', m.group(1))]
        return '(' + '; '.join(author_year(n) for n in nums) + ')'
    t = CITE_RE.sub(repl_standalone, t)
    return t


n_cited = 0
for i, p in enumerate(paras[:i_ref]):
    if CITE_RE.search(p.text):
        new_t = conv_text(p.text)
        n_cited += len(CITE_RE.findall(p.text))
        for r in list(p.runs):
            r.text = ''
        (p.runs[0] if p.runs else p.add_run()).text = new_t
print(f'① 正文引用点转换：{n_cited} 处')

# ---------- ② 文献表：作者-年份格式 + 字母序 ----------
def conv_ref(t):
    """'[n] Surname, I.I., Surname, I., et al.: Title. Venue vol(issue), pages (year). doi' → 作者-年份式"""
    m = re.match(r'^\[(\d+)\]\s+(.*)$', t)
    num, body = int(m.group(1)), m.group(2)
    # 作者段（至第一个 ': '）
    auth, rest = body.split(': ', 1)
    # 作者格式化："Surname, I.I." → "Surname II"；"et al." 去点
    auth = auth.replace(', et al.', ' et al')
    auth = re.sub(r'([A-Z][^,]*?),\s*((?:[A-Z]\.?\s*)+?)(?=,| et al|$)',
                  lambda mm: mm.group(1) + ' ' + mm.group(2).replace('.', '').strip(), auth)
    # 抽年份（最后一个 (YYYY) 或 (YYYYx)）
    ym = list(re.finditer(r'\((\d{4}[ab]?)\)', rest))
    assert ym, f'无年份: {t[:60]}'
    year = ym[-1].group(1)
    rest = rest[:ym[-1].start()].rstrip() + ' ' + rest[ym[-1].end():].strip()
    # 细清理：卷, 页 → 卷: 页；pp. → pp；Proc. → Proc；venue 内 '. ' → ' '
    rest = re.sub(r'pp\.\s*', 'pp ', rest)
    rest = re.sub(r'(\d+\(\d+\)|\d+),\s*(\d+[–-]\d+|\d+)', r'\1:\2', rest)
    rest = rest.replace('Proc. ', 'Proc ').replace('Front. Genet.', 'Front Genet') \
        .replace('Appl. Anim. Behav. Sci.', 'Appl Anim Behav Sci').replace('Anim. Front.', 'Anim Front') \
        .replace('Sci. Rep.', 'Sci Rep').replace('Comput. Electron. Agric.', 'Comput Electron Agric') \
        .replace('Neural Netw.', 'Neural Netw').replace('IEEE Trans. Pattern Anal. Mach. Intell.', 'IEEE Trans Pattern Anal Mach Intell') \
        .replace('J. Real-Time Image Process.', 'J Real-Time Image Process') \
        .replace('INMATEH Agric. Eng.', 'INMATEH Agric Eng') \
        .replace('Smart Agric. Technol.', 'Smart Agric Technol') \
        .replace('Porcine Health Manag.', 'Porcine Health Manag')
    # Luo 两篇同年需 a/b 后缀（与文内 REFKEY 对应）：[33] Collaborative=2025a, [27] PBR-YOLO=2025b
    suffix = {27: 'b', 33: 'a'}.get(num)
    if suffix:
        year = year + suffix
    return f'{auth} ({year}) {rest}'


# 数据集/软件条目手工覆盖（格式与期刊论文不同源）
OVERRIDE = {
    19: 'Roboflow (2026a) Pig behavior dataset, version 1. Roboflow Universe. '
        'https://universe.roboflow.com/km-sd0ce/pig-behavior-wlvku. Accessed 20 July 2026',
    26: 'Roboflow (2026b) Comportamentos dataset. Roboflow Universe. '
        'https://universe.roboflow.com/maria-dnxxx/comportamentos-vdzlw. Accessed 18 August 2026',
    21: 'Jocher G, Qiu J (2024) Ultralytics YOLO11 (v11.0.0) [Computer software]. '
        'https://github.com/ultralytics/ultralytics',
    24: 'Jocher G, Chaurasia A, Qiu J (2023) Ultralytics YOLOv8 (v8.0.0) [Computer software]. '
        'https://github.com/ultralytics/ultralytics',
    23: 'Jocher G, Chaurasia A, Stoken A et al (2022) ultralytics/yolov5: v7.0—YOLOv5 SOTA '
        'realtime instance segmentation. Zenodo. https://doi.org/10.5281/zenodo.7347926',
    14: 'Howard AG, Zhu M, Chen B et al (2017) MobileNets: efficient convolutional neural '
        'networks for mobile vision applications. arXiv preprint arXiv:1704.04861',
    29: 'Yaman B, Mahmud T, Liu CH (2023) Instance-aware repeat factor sampling for '
        'long-tailed object detection. arXiv preprint arXiv:2305.08069',
    30: 'Crasto N (2024) Class imbalance in object detection: an experimental diagnosis '
        'and study of mitigation strategies. arXiv preprint arXiv:2403.07113',
    13: 'Lv J, Wang G, Zhang M et al (2026) WFE-YOLO: a lightweight pig behavior detection '
        'model for livestock farming applications. INMATEH Agric Eng 78(1):1260–1273. '
        'https://doi.org/10.35633/inmateh-78-99',
}

ref_paras = []  # (段落对象, 原编号, 新文本)
for p in paras[i_ref + 1:]:
    t = p.text.strip()
    if re.match(r'^\[\d+\]', t):
        num = int(re.match(r'^\[(\d+)\]', t).group(1))
        ref_paras.append((p, num, OVERRIDE.get(num) or conv_ref(t)))
assert len(ref_paras) == 33, len(ref_paras)

# 字母序排序：物理段落顺序不变，排序后的文本按序写回
paras_in_order = [p for p, _, _ in ref_paras]          # 文档物理顺序
texts_sorted = sorted((t for _, _, t in ref_paras), key=str.lower)
for p, new_t in zip(paras_in_order, texts_sorted):
    for r in list(p.runs):
        r.text = ''
    (p.runs[0] if p.runs else p.add_run()).text = new_t
print('② 文献表 33 条已转作者-年份格式并按字母序重排')

# ---------- ③ 结论并入讨论（TAHP：Discussion 含结论） ----------
i_concl_h = next(i for i, p in enumerate(paras) if p.text.strip().startswith('8. Conclusion'))
i_concl_p = i_concl_h + 1
i_disc = next(i for i, p in enumerate(paras) if p.text.strip().startswith('7. Discussion'))
# 讨论章最后一段 = 结论段前（Limitations 段为讨论末段）
i_decl = next(i for i, p in enumerate(paras) if 'Statements and Declarations' in p.text)
concl_el = paras[i_concl_p]._element
# 把结论段移动到讨论章末尾（Limitations 段之后 = 声明节标题之前）
paras[i_decl]._element.addprevious(concl_el)
# 删除"8. Conclusion"标题段
paras[i_concl_h]._element.getparent().remove(paras[i_concl_h]._element)
print('③ 结论段已并入讨论章末，"8. Conclusion" 标题已删')

# ---------- ④ Statement of Animal Rights（声明节内补充） ----------
paras = doc.paragraphs
i_decl = next(i for i, p in enumerate(paras) if 'Statements and Declarations' in p.text)
decl_par = paras[i_decl + 1]  # Author Contributions 段
new_par = decl_par._element
from docx.oxml.ns import qn
import copy
sar = copy.deepcopy(new_par)
new_par.addnext(sar)
# 找到新段落对象并写文本
para_objs = doc.paragraphs
idx = next(i for i, p in enumerate(para_objs) if p.text.startswith('Author Contributions'))
target = para_objs[idx + 1]
for r in list(target.runs):
    r.text = ''
(target.runs[0] if target.runs else target.add_run()).text = (
    'Statement of Animal Rights: This study used only publicly available, non-invasively '
    'acquired video data recorded by fixed cameras on commercial farms; no animal experiments '
    'were performed by the authors, and no institutional ethics approval was required.')
print('④ Statement of Animal Rights 已加入声明节')

# ---------- ⑤ A4 + 双倍行距 + 连续行号 ----------
for sec in doc.sections:
    sec.page_width, sec.page_height = Mm(210), Mm(297)
    sectPr = sec._sectPr
    if sectPr.find(qn('w:lnNumType')) is None:
        ln = sectPr.makeelement(qn('w:lnNumType'),
                                {qn('w:countBy'): '1', qn('w:restart'): 'continuous'})
        sectPr.append(ln)
for p in doc.paragraphs:
    p.paragraph_format.line_spacing = 2.0
doc.styles['Normal'].font.name = 'Times New Roman'
doc.styles['Normal'].font.size = Pt(10)
print('⑤ A4 + 双倍行距 + 连续行号已设置')

doc.save(DST)

# ---------- 验证 ----------
d2 = Document(DST)
body = '\n'.join(p.text for p in d2.paragraphs[:next(i for i, p in enumerate(d2.paragraphs) if p.text.strip() == 'References')])
assert not CITE_RE.search(body), '正文残留 [n] 引用'
refs = [p.text for p in d2.paragraphs[next(i for i, p in enumerate(d2.paragraphs) if p.text.strip() == 'References') + 1:] if p.text.strip()]
refs = [r for r in refs if re.match(r'^[A-Z][a-z]', r)]
assert len(refs) >= 33, len(refs)
for probe in ['(Canario et al. 2020; Peden et al. 2018)', 'Tu et al. (2022)', '(Roboflow 2026a)',
              'Statement of Animal Rights', '(Luo et al. 2025a)', '(Luo et al. 2025b)']:
    assert probe in body or probe in '\n'.join(refs), f'复检失败: {probe}'
sorted_check = [r.lower() for r in refs[:33]]
assert sorted_check == sorted(sorted_check), '文献表非字母序'
assert '8. Conclusion' not in '\n'.join(p.text for p in d2.paragraphs)
print('验证通过：无 [n] 残留、33 条字母序、结构探针全中 →', DST)
for r in refs[:33]:
    print('  ', r[:95])
