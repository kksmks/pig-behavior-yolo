#!/usr/bin/env python3
"""v5 文献补强手术 —— JRTIP-paper-v4.docx → JRTIP-paper-v5.docx（不动 v4）。

内容（2026-08-07 评估结论，文献均已网络核实）：
  1. 2.1 末尾补 PBR-YOLO [27] + Rahman YOLO11n/TensorRT [28]（2025-2026 时效性）
  2. 2.2 末尾补 IRFS [29] + Crasto 诊断 [30]（采样最近先行工作，差异化表述）
  3. 2.3 末尾补 JRTIP 本刊两篇 FasterNet 先例 [31,32] + 剪枝/蒸馏路线 [33]
     （主动承认 FasterNet+YOLO 不新，立意转向 identity-preserving）
  4. 7 Discussion 首段末尾补与相邻路线的对比段
  5. 4.4 末尾补一句指向新增 Table S2（测试集每类 AP）
  6. Keywords 8 个精简为 6 个
  7. 文末追加参考文献 [27]–[33]
  8. Supplementary：Table S1 补 M5 active 缺失格（0.452，本地 CPU 官方协议复测），
     新增 Table S2（测试集每类 AP，Baseline/M4/M5）

运行：python scripts/apply_v5_edits.py
输出：paper/JRTIP-paper-v5.docx + paper/Supplementary_Material_v5.docx
"""
from pathlib import Path

from docx import Document

PAPER = Path(r'E:/pig-behavior-yolo/paper')
SRC = PAPER / 'JRTIP-paper-v4.docx'
DST = PAPER / 'JRTIP-paper-v5.docx'
SUPP_SRC = PAPER / 'Supplementary_Material.docx'
SUPP_DST = PAPER / 'Supplementary_Material_v5.docx'

# ---------------- 主稿：追加段落文本 ----------------
APPENDS = [
    ("few report per-class results on rare behaviors.",
     " Two very recent studies push the same deployment agenda with different emphases. "
     "PBR-YOLO [27] combines a GhostNet backbone, a FasterNet block with efficient multi-scale "
     "attention, and a lightweight multi-path head for eight piglet behaviors, reporting 15.8 ms "
     "per frame with TensorRT on a Jetson Orin NX—an attention-augmented design demonstrated on a "
     "board far more powerful than the Nano-class hardware we target. Rahman et al. [28] trim "
     "YOLO11n for the single posture–nursing task in farrowing crates (removing the small-object "
     "detection head) and accelerate it with TensorRT; class imbalance, which is central to our "
     "setting, does not arise in their formulation. Neither reports per-class accuracy on rare "
     "behaviors, and both evaluate on random frame splits."),

    ("an aspect rarely ablated in this literature.",
     " Instance-aware repeat factor sampling [29] refines RFS by counting annotated instances; "
     "we instead count class-bearing images—the unit that duplication actually replicates—and add "
     "a square-root damper with a hard cap so that a handful of rare scenes cannot dominate an "
     "epoch. A systematic diagnosis of mitigation strategies for single-stage detectors [30] "
     "likewise finds no universally dominant family, which supports our choice of the least "
     "intrusive one."),

    ("We document both explicitly in Section 5.",
     " FasterNet-based substitution is itself established—Guo et al. [31] integrated FasterNet "
     "into YOLOv8 for real-time underwater detection and YOLO-FGD [32] applied the same idea to "
     "PCB inspection, both in this journal—so our contribution is not the substitution but the "
     "condition under which it preserves accuracy: identity-preserving weight transfer "
     "(Section 3.5), which these works do not discuss. A complementary route to lightness is "
     "post-training compression: Luo et al. [33] combine channel pruning with knowledge "
     "distillation for piglet behavior recognition, accepting a small accuracy margin in exchange "
     "for large parameter cuts without changing the training-time architecture."),

    ("this is a data-scarcity problem, not an architecture problem.",
     " Per-class AP50 on the held-out test set for the baseline, M4 and M5 is reported in "
     "Table S2 of the Supplementary Material; the rare-class pattern mirrors the validation set."),

    ("readers should expect the same on similar boards.",
     " The comparison with neighboring approaches is instructive. Attention-augmented detectors "
     "such as PBR-YOLO [27] report strong results with attention-augmented lightweight designs, "
     "but our two failed integrations (M1, M2) show that these gains do not transfer when the "
     "modules are inserted into an already pre-trained pathway—the identity-preserving route is "
     "what makes backbone substitution practical in transfer-heavy pipelines. Post-training "
     "compression [33] is orthogonal to both our data-level and backbone-level changes, and the "
     "two routes could in principle be stacked; we leave that combination to future work."),
]

KEYWORDS_NEW = ("Keywords: pig behavior detection; class imbalance; FasterNet; edge computing; "
                "Jetson Nano; precision livestock farming")

NEW_REFS = [
    "[27] Luo, Y., Lin, K., Xiao, Z., et al.: PBR-YOLO: a lightweight piglet multi-behavior "
    "recognition algorithm based on improved YOLOv8. Smart Agric. Technol. 10, 100785 (2025). "
    "https://doi.org/10.1016/j.atech.2025.100785",

    "[28] Rahman, M., Souza, V.H.S., Brown-Brandl, T.M., et al.: Accelerating sow nursing "
    "behavior monitoring with modified YOLO11n architecture and TensorRT integration. "
    "Porcine Health Manag. (2026). https://doi.org/10.1186/s40813-026-00507-3",

    "[29] Yaman, B., Mahmud, T., Liu, C.H.: Instance-aware repeat factor sampling for long-tailed "
    "object detection. arXiv preprint arXiv:2305.08069 (2023)",

    "[30] Crasto, N.: Class imbalance in object detection: an experimental diagnosis and study of "
    "mitigation strategies. arXiv preprint arXiv:2403.07113 (2024)",

    "[31] Guo, A., Sun, K., Zhang, Z.: A lightweight YOLOv8 integrating FasterNet for real-time "
    "underwater object detection. J. Real-Time Image Process. 21(2), 49 (2024). "
    "https://doi.org/10.1007/s11554-024-01431-x",

    "[32] Qin, C., Zhou, Z.: YOLO-FGD: a fast lightweight PCB defect method based on FasterNet "
    "and the Gather-and-Distribute mechanism. J. Real-Time Image Process. 21(4), 122 (2024). "
    "https://doi.org/10.1007/s11554-024-01504-x",

    "[33] Luo, Y., Lin, K., Xiao, Z., et al.: Collaborative optimization of model pruning and "
    "knowledge distillation for efficient and lightweight multi-behavior recognition in piglets. "
    "Animals 15(11), 1563 (2025). https://doi.org/10.3390/ani15111563",
]

# ---------------- 补充材料：Table S2 数据（官方协议，2026-08-07 本地复核） ----------------
S2_ROWS = [
    ("active",        "0.526", "0.639", "0.533"),
    ("drink",         "0.389", "0.373", "0.408"),
    ("eat",           "0.445", "0.418", "0.429"),
    ("fight",         "0.840", "0.871", "0.849"),
    ("investigating", "0.552", "0.562", "0.588"),
    ("lying",         "0.741", "0.737", "0.722"),
    ("nose-to-nose",  "0.783", "0.740", "0.749"),
    ("sitting",       "0.534", "0.571", "0.568"),
    ("standing",      "0.487", "0.468", "0.475"),
    ("walk",          "0.667", "0.655", "0.612"),
    ("mAP50 (all)",   "0.596", "0.604", "0.593"),
]
S2_CAPTION = ("Table S2. Per-class AP50 on the held-out test set for the baseline and the two "
              "adopted variants. The rare-class pattern mirrors the validation set (Table S1): "
              "the gains of M4 concentrate on active, sitting and fight; M5 keeps part of the "
              "rare-class gain while carrying 4.4% fewer parameters. M5 values were re-measured "
              "under the official validation protocol (overall mAP50 0.5933, matching the "
              "cloud-reported 0.5932).")


def edit_main():
    doc = Document(str(SRC))
    done = set()

    for para in doc.paragraphs:
        t = para.text
        if not t.strip():
            continue
        # 关键词整段替换
        if t.startswith('Keywords:'):
            para.runs[0].text = KEYWORDS_NEW
            for r in para.runs[1:]:
                r.text = ''
            done.add('keywords')
            continue
        # 锚点追加
        for i, (anchor, addition) in enumerate(APPENDS):
            if t.rstrip().endswith(anchor) and i not in done:
                para.add_run(addition)
                done.add(i)
                break

    # 参考文献追加（[26] 为当前最后一条）
    ref_start = len(doc.paragraphs)
    for ref in NEW_REFS:
        doc.add_paragraph(ref)
    done.add('refs')

    doc.save(str(DST))
    print(f'[OK] 主稿 → {DST}')
    print(f'  追加段落 {len([d for d in done if isinstance(d, int)])}/5，关键词 {"✓" if "keywords" in done else "✗"}，参考文献 +{len(NEW_REFS)} 条')
    missing = [APPENDS[i][0][:50] for i in range(len(APPENDS)) if i not in done]
    if missing:
        print('  [WARN] 未命中锚点:', missing)
    return len(missing) == 0


def edit_supp():
    doc = Document(str(SUPP_SRC))
    # 1) Table S1 补 M5 active
    t = doc.tables[0]
    header = [c.text.strip() for c in t.rows[0].cells]
    m5_col = header.index('M5 (combined)')
    filled = False
    for row in t.rows[1:]:
        if row.cells[0].text.strip() == 'active':
            assert row.cells[m5_col].text.strip() in ('—', '-', ''), row.cells[m5_col].text
            row.cells[m5_col].text = '0.452'
            filled = True
    print(f'  Table S1 M5 active 补格 {"✓ (0.452)" if filled else "✗ 未找到"}')

    # 2) 追加 Table S2
    cap = doc.add_paragraph(S2_CAPTION)
    cap.runs[0].font.size = None  # 跟随正文样式
    table = doc.add_table(rows=1 + len(S2_ROWS), cols=4)
    try:
        table.style = t.style  # 与 S1 同款
    except Exception:
        pass
    for j, h in enumerate(['Class', 'Baseline', 'M4 (sampling)', 'M5 (combined)']):
        table.rows[0].cells[j].text = h
    for i, row in enumerate(S2_ROWS):
        for j, v in enumerate(row):
            table.rows[i + 1].cells[j].text = v

    doc.save(str(SUPP_DST))
    print(f'[OK] 补充材料 → {SUPP_DST}（Table S2 已追加，{len(S2_ROWS)} 行）')
    return filled


if __name__ == '__main__':
    ok1 = edit_main()
    ok2 = edit_supp()
    print('全部完成' if (ok1 and ok2) else '有未命中项，请检查上方 WARN')
