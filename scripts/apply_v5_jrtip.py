# -*- coding: utf-8 -*-
"""JRTIP submission-guideline compliance fixes for v5 (2026-08-07).

Guideline sources: springer.com/journal/11554/submission-guidelines
  1. Abstract: no undefined abbreviations -> expand mAP50 and FPS at first use.
  2. Figure captions: no punctuation at the end -> strip trailing '.' (Fig. 5, Fig. 8).
  3. Declarations heading -> "Statements and Declarations".
  4. SI must be cited as "Online Resource 1" (not "the Supplementary Material").
  5. SI file must carry journal name + author block -> add header line.
"""
from docx import Document

PAPER = r"E:\pig-behavior-yolo\paper\JRTIP-paper-v5.docx"
SUPP = r"E:\pig-behavior-yolo\paper\Supplementary_Material_v5.docx"

PAPER_REPL = [
    (
        "(test mAP50 0.590 ± 0.009 vs. a three-seed baseline of 0.606 ± 0.008)",
        "(test mean average precision at IoU 0.5 (mAP50) of 0.590 ± 0.009 vs. "
        "a three-seed baseline of 0.606 ± 0.008)",
    ),
    (
        "(117.6 FPS on an RTX 3090)",
        "(117.6 frames per second (FPS) on an RTX 3090)",
    ),
    ("Table S1 in the Supplementary Material", "Table S1 in Online Resource 1"),
    ("Table S2 of the Supplementary Material", "Table S2 in Online Resource 1"),
]

hits = {}


def replace_in_runs(paras, old, new, key):
    for p in paras:
        for r in p.runs:
            if old in r.text:
                r.text = r.text.replace(old, new)
                hits[key] = hits.get(key, 0) + 1
                return


doc = Document(PAPER)
for i, (old, new) in enumerate(PAPER_REPL):
    replace_in_runs(doc.paragraphs, old, new, f"paper-{i}")

# Declarations heading
for p in doc.paragraphs:
    if p.text.strip() == "Declarations":
        p.runs[0].text = p.runs[0].text.replace("Declarations", "Statements and Declarations")
        for r in p.runs[1:]:
            r.text = ""
        hits["decl-heading"] = 1
        break

# Figure caption trailing periods (Fig. 5 / Fig. 8)
hits["fig-period"] = 0
for p in doc.paragraphs:
    s = p.text.strip()
    if (s.startswith("Fig. 5") or s.startswith("Fig. 8")) and s.endswith("."):
        for r in reversed(p.runs):
            if r.text.rstrip().endswith("."):
                idx = r.text.rstrip().rfind(".")
                keep_trailing_ws = r.text[len(r.text.rstrip()):]
                r.text = r.text.rstrip()[:idx] + keep_trailing_ws
                hits["fig-period"] += 1
                break

ok = (all(hits.get(f"paper-{i}", 0) == 1 for i in range(len(PAPER_REPL)))
      and hits.get("decl-heading") == 1 and hits.get("fig-period") == 2)
print("paper hits:", {k: v for k, v in hits.items() if k != "fig-period"}, "fig-period:", hits["fig-period"])
if ok:
    doc.save(PAPER)
    print("saved:", PAPER)

# --- Supplementary header ---
sd = Document(SUPP)
p0 = sd.paragraphs[0]
if "Online Resource 1" not in p0.text:
    p0.runs[0].text = p0.runs[0].text.replace("Supplementary Material",
                                              "Supplementary Material (Online Resource 1)")
anchor = sd.paragraphs[2]  # first Table caption
anchor.insert_paragraph_before(
    "Journal of Real-Time Image Processing · [Author names] · Corresponding author: [e-mail]")
sd.save(SUPP)
print("saved:", SUPP)

# --- verify abstract word count ---
d2 = Document(PAPER)
paras2 = [p.text for p in d2.paragraphs]
abstract = paras2[paras2.index("Abstract") + 1]
print("摘要词数(改后):", len(abstract.split()))
print("图5/8题注尾字符:", [p.text.strip()[-1] for p in d2.paragraphs
      if p.text.strip().startswith(("Fig. 5", "Fig. 8"))])
