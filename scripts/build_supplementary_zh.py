# 生成中文补充材料 DOCX（表 S1：验证集分类别 AP50）——与 Supplementary_Material.docx 对齐
# 运行：python scripts/build_supplementary_zh.py → paper/猪行为检测-补充材料-v4.docx
import sys
sys.stdout.reconfigure(encoding='utf-8', errors='replace')
from docx import Document
from docx.shared import Pt

doc = Document()
doc.styles['Normal'].font.name = 'SimSun'
doc.styles['Normal'].font.size = Pt(11)

doc.add_heading('补充材料', level=1)
doc.add_paragraph('面向边缘设备的群养猪多行为实时检测：类别不均衡感知采样与轻量化 FasterNet 主干 —— 补充材料')

cap = doc.add_paragraph('表 S1 验证集分类别 AP50。独立测试集上的关键稀有类结果：'
                        'active 0.526→0.639（基线→M4，+11.3 点）；sitting 0.534→0.571；fight 0.840→0.871。')
cap.runs[0].font.size = Pt(9)

headers = ['类别', '基线', 'M3（FasterNet）', 'M4（采样）', 'M5（组合）', 'M4−基线']
rows = [
    ['active', '0.459', '0.372', '0.552', '—', '+9.3'],
    ['drink', '0.408', '0.440', '0.459', '0.430', '+5.1'],
    ['eat', '0.436', '0.437', '0.447', '0.449', '+1.1'],
    ['fight', '0.858', '0.828', '0.842', '0.786', '-1.6'],
    ['investigating', '0.581', '0.571', '0.570', '0.577', '-1.1'],
    ['lying', '0.783', '0.764', '0.766', '0.743', '-1.7'],
    ['nose-to-nose', '0.686', '0.655', '0.622', '0.710', '-6.4'],
    ['sitting', '0.403', '0.233', '0.423', '0.349', '+2.0'],
    ['standing', '0.438', '0.429', '0.449', '0.445', '+1.1'],
    ['walk', '0.687', '0.646', '0.686', '0.667', '-0.1'],
    ['mAP50（全部）', '0.573', '0.537', '0.582', '0.561', '+0.9'],
]
t = doc.add_table(rows=1 + len(rows), cols=len(headers))
t.style = 'Light Grid Accent 1'
for j, htxt in enumerate(headers):
    t.rows[0].cells[j].text = htxt
for i, row in enumerate(rows):
    for j, v in enumerate(row):
        t.rows[i + 1].cells[j].text = v

out = 'paper/猪行为检测-补充材料-v4.docx'
doc.save(out)
print('已生成:', out)
