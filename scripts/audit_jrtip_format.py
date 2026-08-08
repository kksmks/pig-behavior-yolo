# -*- coding: utf-8 -*-
"""
audit_jrtip_format.py — JRTIP 官方投稿指南合规体检（对 paper/JRTIP-paper-v5.docx 等）。
依据：https://link.springer.com/journal/11554/submission-guidelines（2026-08-08 检索）
输出：逐项 PASS / FAIL / WARN / MANUAL（需人工或用户填写）。
只读不改。
"""
import re
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent.parent
MS = ROOT / 'paper' / 'JRTIP-paper-v5.docx'
SI = ROOT / 'paper' / 'Supplementary_Material_v5.docx'
CL = ROOT / 'paper' / 'cover-letter-jrtip.docx'

results = []

def rep(item, status, detail):
    results.append((item, status, detail))
    print(f'[{status:6s}] {item}: {detail}')

doc = Document(str(MS))
paras = doc.paragraphs
texts = [p.text for p in paras]
full = '\n'.join(texts)

# ---------- 1. 标题页 ----------
title = next((t for t in texts if t.strip().startswith('Real-Time')), '')
rep('标题简洁 informative', 'PASS' if title else 'FAIL', title[:80] or '未找到')

author_line = texts[1] if len(texts) > 1 else ''
if 'to be filled' in author_line or 'Authors' in author_line:
    rep('作者/单位/通讯邮箱/ORCID', 'MANUAL', '占位行未填（用户必办）: ' + author_line[:60])
else:
    rep('作者/单位/通讯邮箱/ORCID', 'PASS', author_line[:70])

# ---------- 2. 摘要 150-250 词 ----------
abs_idx = next(i for i, p in enumerate(paras) if p.style.name == 'Heading 1' and p.text.strip() == 'Abstract')
kw_idx = next(i for i, t in enumerate(texts) if t.strip().lower().startswith('keywords'))
abstract = ' '.join(texts[abs_idx + 1:kw_idx])
nw = len(abstract.split())
rep('摘要词数 150–250', 'PASS' if 150 <= nw <= 250 else 'FAIL', f'{nw} 词')
rep('摘要无引用标注', 'PASS' if not re.search(r'\[\d+\]', abstract) else 'FAIL',
    '无 [n]' if not re.search(r'\[\d+\]', abstract) else '含引用！')

# ---------- 3. 关键词 4-6 ----------
kw_text = texts[kw_idx]
kws = [k for k in re.split(r'[;；]', kw_text.split(':', 1)[-1]) if k.strip()]
rep('关键词 4–6 个', 'PASS' if 4 <= len(kws) <= 6 else 'FAIL', f'{len(kws)} 个: {kw_text[:90]}')

# ---------- 4. 标题十进制、层级 ≤3 ----------
heads = [(p.style.name, p.text.strip()) for p in paras if p.style.name.startswith('Heading')]
bad_lvl = [h for h in heads if h[0] == 'Heading 4']
sec_nums = [h[1] for h in heads if re.match(r'^\d', h[1])]
lvl3 = [h for h in heads if h[0] == 'Heading 3']
rep('标题十进制编号', 'PASS' if sec_nums else 'FAIL', f'{len(sec_nums)} 个编号节标题')
rep('层级不超过 3 级', 'PASS' if not bad_lvl else 'FAIL',
    f'最深 Heading {max(int(s.split()[-1]) for s, _ in heads)} 级' if heads else '?')

# ---------- 5. 图：编号、顺序引用、题注格式 ----------
body_before_refs = full.split('\nReferences\n')[0] if '\nReferences\n' in full else full
fig_cite_seq = [int(m) for m in re.findall(r'Fig\.\s*(\d+)', body_before_refs)]
first_cites = []
seen = set()
for n in fig_cite_seq:
    if n not in seen:
        seen.add(n)
        first_cites.append(n)
rep('图按顺序引用', 'PASS' if first_cites == sorted(first_cites) == list(range(1, 9)) else 'WARN',
    f'首现顺序 {first_cites}')

fig_caps = [t.strip() for t in texts if re.match(r'^Fig\.\s*\d', t.strip())]
bad_caps = [c for c in fig_caps if c.rstrip().endswith(('.', '。'))]
rep('图题注尾无句号（Springer 规则）', 'PASS' if not bad_caps else 'FAIL',
    f'{len(fig_caps)} 条图题注，违规 {len(bad_caps)}')
subfig = [c for c in fig_caps if re.search(r'\([a-z]\)', c)]
rep('分图小写 (a)(b) 标注', 'PASS' if subfig else 'WARN', f'{len(subfig)} 条含分图标注')

# ---------- 6. 表：编号、顺序引用 ----------
tab_cite_seq = [int(m) for m in re.findall(r'Table\s+(\d+)', body_before_refs)]
seen_t, first_t = set(), []
for n in tab_cite_seq:
    if n not in seen_t:
        seen_t.add(n)
        first_t.append(n)
rep('表按顺序引用', 'PASS' if first_t == sorted(first_t) else 'WARN', f'首现顺序 {first_t}')
rep('表格数量', 'PASS' if len(doc.tables) == 7 else 'WARN', f'{len(doc.tables)} 张（预期 7）')

# ---------- 7. 参考文献 ----------
ref_start = next(i for i, p in enumerate(paras) if p.style.name == 'Heading 1' and p.text.strip() == 'References')
ref_entries = [t.strip() for t in texts[ref_start + 1:] if re.match(r'^\[\d+\]', t.strip())]
nums = [int(re.match(r'^\[(\d+)\]', t).group(1)) for t in ref_entries]
rep('文献编号连续 [1]–[n]', 'PASS' if nums == list(range(1, len(nums) + 1)) else 'FAIL',
    f'{len(nums)} 条, max [{max(nums) if nums else 0}]')

cited = set()
for grp in re.findall(r'\[([\d,\s\-–]+)\]', body_before_refs):
    for part in grp.split(','):
        part = part.strip().replace('–', '-')
        if '-' in part:
            a, _, b = part.partition('-')
            if a.strip().isdigit() and b.strip().isdigit():
                cited.update(range(int(a), int(b) + 1))
        elif part.isdigit():
            cited.add(int(part))
cited = {n for n in cited if n <= len(nums)}
uncited = set(nums) - cited
rep('全部文献被正文引用', 'PASS' if not uncited else 'FAIL', f'未被引: {sorted(uncited) or "无"}')

no_doi = [t[:60] for t in ref_entries if 'doi.org' not in t.lower()]
rep('DOI 全链接 (https://doi.org/)', 'WARN' if no_doi else 'PASS',
    f'{len(no_doi)} 条无 DOI 链接（书籍/会议可豁免）' if no_doi else '全部含 DOI')

# ---------- 8. Statements and Declarations ----------
sd_idx = next((i for i, p in enumerate(paras) if p.style.name == 'Heading 1' and 'Statements and Declarations' in p.text), None)
if sd_idx is None:
    rep('Statements and Declarations 节', 'FAIL', '未找到')
else:
    sd_text = '\n'.join(texts[sd_idx:ref_start])
    for sub in ['Funding', 'Competing Interests', 'Data Availability', 'Author Contribution']:
        if sub == 'Competing Interests':
            found = re.search(r'Competing Interests|Conflicts? of Interest', sd_text, re.I)
        else:
            found = re.search(re.escape(sub), sd_text, re.I)
        if found:
            placeholder = 'to be filled' in sd_text.lower() or '[to' in sd_text.lower()
            rep(f'  └ {sub}', 'MANUAL' if placeholder else 'PASS',
                '存在（含待填占位）' if placeholder else '存在')
        else:
            rep(f'  └ {sub}', 'FAIL', '缺失')
    rep('  └ AI-Assisted 声明（Springer LLM 政策）',
        'PASS' if re.search(r'AI[- ]assisted|LLM|artificial intelligence', sd_text, re.I) else 'WARN',
        '已声明' if re.search(r'AI[- ]assisted', sd_text, re.I) else '未见明确声明')

# ---------- 9. 伦理合规 ----------
rep('动物福利/IRB 声明', 'PASS' if re.search(r'(animal|IRB|ethic|welfare|consent)', full, re.I) else 'FAIL',
    '含相关声明' if re.search(r'(animal welfare|IRB|Ethics)', full, re.I) else '弱')

# ---------- 10. 脚注/尾注 ----------
import zipfile
z = zipfile.ZipFile(str(MS))
has_fn = 'word/footnotes.xml' in z.namelist() and b'footnoteReference' in z.read('word/footnotes.xml')
has_en = 'word/endnotes.xml' in z.namelist() and b'endnoteReference' in z.read('word/endnotes.xml')
rep('用脚注而非尾注', 'PASS' if not has_en else 'FAIL', f'脚注={has_fn} 尾注={has_en}')

# ---------- 11. SI 引用 "Online Resource n" ----------
rep('SI 引用为 "Online Resource 1"', 'PASS' if 'Online Resource 1' in full else 'FAIL',
    f"出现 {full.count('Online Resource 1')} 次")

# ---------- 12. Data Availability 含仓库链接 ----------
rep('Data Availability 含 GitHub 链接', 'PASS' if 'github.com/kksmks' in full else 'FAIL',
    'github.com/kksmks/pig-behavior-yolo' if 'github.com/kksmks' in full else '缺失')

# ---------- 13. SI 文件头 ----------
si = Document(str(SI))
si_head = '\n'.join(p.text for p in si.paragraphs[:6])
rep('SI 含期刊名+文章标题', 'PASS' if re.search(r'Real-Time Image Processing|Journal', si_head, re.I) else 'WARN',
    si_head.replace('\n', ' | ')[:100])

# ---------- 14. Cover letter real-time ----------
cl = Document(str(CL))
cl_text = '\n'.join(p.text for p in cl.paragraphs)
rep('Cover Letter 论述 real-time 问题', 'PASS' if re.search(r'real[- ]time', cl_text, re.I) else 'FAIL',
    f'real-time 出现 {len(re.findall(chr(114)+chr(101)+chr(97)+chr(108)+chr(45)+chr(116)+chr(105)+chr(109)+chr(101), cl_text, re.I))} 次')

print()
n_pass = sum(1 for _, s, _ in results if s == 'PASS')
n_fail = sum(1 for _, s, _ in results if s == 'FAIL')
n_warn = sum(1 for _, s, _ in results if s == 'WARN')
n_man = sum(1 for _, s, _ in results if s == 'MANUAL')
print(f'== 合计: PASS {n_pass} / FAIL {n_fail} / WARN {n_warn} / MANUAL {n_man} ==')
