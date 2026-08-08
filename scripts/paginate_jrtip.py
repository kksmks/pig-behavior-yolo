# -*- coding: utf-8 -*-
"""
paginate_jrtip.py — 把 JRTIP-paper-v4.docx 灌入 JRTIP 官方 Word 模板的双栏期刊格式，
用于实测排版页数（12 页上限含参考文献与 bio）。

做法：复制 v4 → 套用模板几何（Letter / T17.8mm B17.8mm L16.5mm R16.5mm / 双栏间距 288twips）
+ 模板字号体系（正文 TNR 10pt 单倍行距、标题 24pt、作者 11pt、摘要与关键词 9pt、
图题表题与表格 8pt、参考文献 8pt、H1/H2 10pt 加粗）+ 图片按栏宽重排
（Fig.2/5/8 跨双栏 174mm，其余单栏 84mm，跨栏图用连续分节符实现）。
表格宽度压到栏宽（tblW 100% + autofit），单元格字号 8pt。

产物：paper/template/v4-jrtip-format.docx（供 Word COM 数页数 / 导出 PDF）
纯测量用途，不改 v4 正稿。
"""
import copy
import re
import shutil
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Mm

ROOT = Path(__file__).resolve().parent.parent
SRC = ROOT / 'paper' / 'JRTIP-paper-v5.docx'
TEMPLATE = ROOT / 'paper' / 'template' / 'JRTIP-WordTemplate.docx'
OUT = ROOT / 'paper' / 'template' / 'v5-jrtip-format.docx'

FULLWIDTH_FIGS = {2, 5}             # 跨双栏的图（Fig.8 已定为单栏 362dpi，勿通栏）
COLW_EMU = int(84 * 36000)          # 单栏图宽 84mm
FULLW_EMU = int(174 * 36000)        # 跨栏图宽 174mm


def clone_pgsetup(tpl_doc):
    """取模板第一节的 pgSz/pgMar/cols 三个元素，作为几何真源。"""
    sectPr = tpl_doc.sections[0]._sectPr
    return (copy.deepcopy(sectPr.find(qn('w:pgSz'))),
            copy.deepcopy(sectPr.find(qn('w:pgMar'))),
            copy.deepcopy(sectPr.find(qn('w:cols'))))


def apply_geometry(sectPr, pgSz, pgMar, cols, num_cols):
    for tag, el in (('w:pgSz', pgSz), ('w:pgMar', pgMar), ('w:cols', cols)):
        old = sectPr.find(qn(tag))
        if old is not None:
            sectPr.remove(old)
        new = copy.deepcopy(el)
        if tag == 'w:cols':
            new.set(qn('w:num'), str(num_cols))
        sectPr.append(new)


def make_break_para(pgSz, pgMar, cols, num_cols):
    """生成一个带 sectPr 的空段（连续分节符，终结前面那一节）。"""
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    sectPr = OxmlElement('w:sectPr')
    typ = OxmlElement('w:type')
    typ.set(qn('w:val'), 'continuous')
    sectPr.append(typ)
    apply_geometry(sectPr, pgSz, pgMar, cols, num_cols)
    pPr.append(sectPr)
    p.append(pPr)
    return p


def set_style(st, size_pt, bold=None, center=False, before=None, after=None):
    st.font.name = 'Times New Roman'
    st.font.size = Pt(size_pt)
    if bold is not None:
        st.font.bold = bold
    from docx.oxml.ns import qn as _qn
    rpr = st.element.get_or_add_rPr()
    color = rpr.find(_qn('w:color'))
    if color is None:
        color = OxmlElement('w:color')
        rpr.append(color)
    color.set(_qn('w:val'), '000000')
    pf = st.paragraph_format
    if center:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if before is not None:
        pf.space_before = Pt(before)
    if after is not None:
        pf.space_after = Pt(after)


def strip_run_sizes(doc):
    """剥掉所有 run 级直接字号，让样式接管。"""
    for rpr in doc.element.body.iter(qn('w:rPr')):
        for tag in ('w:sz', 'w:szCs'):
            el = rpr.find(qn(tag))
            if el is not None:
                rpr.remove(el)


def set_runs_size(para, pt_val):
    half = str(int(pt_val * 2))
    for r in para.runs:
        rpr = r._element.get_or_add_rPr()
        for tag in ('w:sz', 'w:szCs'):
            el = rpr.find(qn(tag))
            if el is None:
                el = OxmlElement(tag)
                rpr.append(el)
            el.set(qn('w:val'), half)


def resize_image(run, target_cx):
    for ext in run._element.findall('.//' + qn('wp:extent')):
        cx, cy = int(ext.get('cx')), int(ext.get('cy'))
        ext.set('cx', str(target_cx))
        ext.set('cy', str(int(cy * target_cx / cx)))
    for ext in run._element.findall('.//' + qn('a:ext')):
        cx, cy = int(ext.get('cx')), int(ext.get('cy'))
        ext.set('cx', str(target_cx))
        ext.set('cy', str(int(cy * target_cx / cx)))


def main():
    tpl = Document(str(TEMPLATE))
    pgSz, pgMar, cols = clone_pgsetup(tpl)

    shutil.copy(str(SRC), str(OUT))
    doc = Document(str(OUT))

    # 1) 样式体系对齐模板
    set_style(doc.styles['Normal'], 10)
    doc.styles['Normal'].paragraph_format.space_after = Pt(0)
    set_style(doc.styles['Title'], 24, center=True)
    set_style(doc.styles['Heading 1'], 10, bold=True, before=12, after=4)
    set_style(doc.styles['Heading 2'], 10, bold=True, before=6, after=3)
    strip_run_sizes(doc)

    # 2) 逐段处理：作者行 / 摘要 / 图题表题 / 参考文献 / 图片
    paras = doc.paragraphs
    in_abstract = False
    in_refs = False
    fig_idx = 0
    cap_re = re.compile(r'^(Fig|Table)\.?\s*\d')
    fig_para_info = []           # (fig_no, 图片所在段元素, 图题段元素)

    for i, p in enumerate(paras):
        text = p.text.strip()
        style = p.style.name

        if style == 'Heading 1' and text == 'Abstract':
            in_abstract = True
            continue
        if style == 'Heading 1' and text.startswith('1.'):
            in_abstract = False
        if style == 'Heading 1' and text.startswith('References'):
            in_refs = True
            continue

        if i == 1:               # 作者占位行
            set_runs_size(p, 11)
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue
        if in_abstract and style == 'Normal':
            set_runs_size(p, 9)  # 摘要正文 + Keywords
            continue
        if in_refs and style == 'Normal':
            set_runs_size(p, 8)  # 参考文献条目
            continue
        if cap_re.match(text) and style == 'Normal':
            set_runs_size(p, 8)  # 图题/表题
            continue

        # 图片段
        img_runs = [r for r in p.runs
                    if r._element.findall('.//' + qn('a:blip'))]
        if img_runs:
            fig_idx += 1
            target = FULLW_EMU if fig_idx in FULLWIDTH_FIGS else COLW_EMU
            for r in img_runs:
                resize_image(r, target)
            fig_para_info.append((fig_idx, p._p))

    # 3) 跨栏图包连续分节符（图段+紧随的图题段一起进单栏节）
    for fig_no, p_el in fig_para_info:
        if fig_no not in FULLWIDTH_FIGS:
            continue
        # 图题段 = 图片段的下一个兄弟段
        cap_el = p_el.getnext()
        end_anchor = cap_el if cap_el is not None and cap_el.tag == qn('w:p') else p_el
        p_el.addprevious(make_break_para(pgSz, pgMar, cols, 2))
        end_anchor.addnext(make_break_para(pgSz, pgMar, cols, 1))

    # 4) 表格：栏宽 100% + autofit + 8pt
    for tbl in doc.tables:
        tblPr = tbl._tbl.tblPr
        tblW = tblPr.find(qn('w:tblW'))
        if tblW is None:
            tblW = OxmlElement('w:tblW')
            tblPr.append(tblW)
        tblW.set(qn('w:w'), '5000')
        tblW.set(qn('w:type'), 'pct')
        layout = tblPr.find(qn('w:tblLayout'))
        if layout is None:
            layout = OxmlElement('w:tblLayout')
            tblPr.append(layout)
        layout.set(qn('w:type'), 'autofit')
        for tc in tbl._tbl.iter(qn('w:tc')):
            tcW = tc.find(qn('w:tcPr') + '/' + qn('w:tcW'))
            if tcW is not None:
                tcW.getparent().remove(tcW)
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    set_runs_size(p, 8)

    # 5) 文档级 sectPr → 模板几何 + 双栏
    apply_geometry(doc.sections[-1]._sectPr, pgSz, pgMar, cols, 2)

    doc.save(str(OUT))
    print(f'OK -> {OUT}  (figures: {fig_idx}, fullwidth: {sorted(FULLWIDTH_FIGS)})')


if __name__ == '__main__':
    main()
